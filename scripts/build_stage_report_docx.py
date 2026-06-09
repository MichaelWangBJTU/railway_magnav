from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROC_DIR = Path(r"C:\Users\m1352\Desktop\磁导航\数据\codex_railway_magnav\data_proc_new")
OUT_DOCX = PROC_DIR / "铁路地磁导航数据处理阶段性汇报.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)


def style_table(table, header_fill: str = "E8EEF5") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.name = "Microsoft YaHei"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    r.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        p = doc.add_paragraph(f"[缺少图片：{path.name}]")
        p.style = "Body Text"
        return
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("铁路地磁导航数据处理阶段性汇报")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 77, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("良陈铁路约 700 m 区段 | SPAN GPGGA + 三轴磁强计")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(89, 89, 89)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], title: str | None = None) -> None:
    if title:
        doc.add_heading(title, level=3)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "项目", True)
    set_cell_text(hdr[1], "结果", True)
    for key, val in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key)
        cells[1].text = val
        for p in cells[1].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = "Microsoft YaHei"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                r.font.size = Pt(9)
    style_table(table)


def add_dataframe_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], max_rows: int | None = None) -> None:
    if max_rows:
        df = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = row[col]
            if isinstance(val, float):
                txt = f"{val:.3g}"
            else:
                txt = str(val)
            set_cell_text(cells[i], txt)
    style_table(table)


def main() -> None:
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("1. 汇报摘要", level=1)
    add_bullets(
        doc,
        [
            "已完成 SPAN GPGGA 与磁强计数据的时间对齐、统一轨道坐标轴建图，以及 4.14/5.13 两次采集的磁图对比。",
            "GPGGA 时间按 NovAtel/Hexagon OEM7 文档中的 UTC 字段处理，转换为北京时间时只加 8 小时，不减闰秒。",
            "已确认小车往返掉头会改变车体系 xyz 方向；当前输出中保留 raw xyz，同时新增每趟去基线后的轨道坐标系磁异常分量。",
            "现阶段最可靠的匹配特征仍是总场 total；Y anomaly 和 X anomaly 有一定跨日期相似性，但三轴分量还需要更系统的标定与补偿。",
        ],
    )

    add_key_value_table(
        doc,
        [
            ("数据日期", "2026-04-13 与 2026-05-13"),
            ("空间网格", "统一轨道 0 点，沿轨道方向 0.5 m 间隔"),
            ("建图坐标", "4.14 与 5.13 共用同一条轨道轴和同一物理 0 点"),
            ("磁图融合", "每个距离点对多趟观测取稳健中位数；同时保留均值、标准差和 MAD"),
            ("输出目录", str(PROC_DIR)),
        ],
        "关键配置",
    )

    doc.add_heading("2. 数据处理流程", level=1)
    add_numbered(
        doc,
        [
            "读取 SPAN 转换后的 GPGGA 位置数据，并将 GPGGA UTC 时间转换为北京时间。",
            "读取磁强计数据，解析系统北京时间、Mag_X/Mag_Y/Mag_Z、Pitch/Roll/Yaw。",
            "用全部有效 GPGGA 坐标拟合铁路一维方向轴，定义统一绝对距离 s。",
            "按时间将磁强计样本插值匹配到 SPAN 坐标，得到每个磁样本对应的经纬度和 s。",
            "根据 s 随时间增减判断小车方向；对每趟先在车体系下去除中位数偏置，再转换到统一轨道坐标系。",
            "将各趟数据插值到 0.5 m 地图点，生成 4.14 和 5.13 的融合磁图与对比图。",
        ],
    )

    doc.add_heading("3. 方向修正与问题定位", level=1)
    p = doc.add_paragraph()
    p.add_run("传感器坐标定义：").bold = True
    p.add_run("X 轴指向小车正后方，Y 轴垂直地面向下，Z 轴指向小车正右方。因此小车掉头后，车体系 X/Z 的物理方向会改变。")
    p = doc.add_paragraph()
    p.add_run("当前采用的修正：").bold = True
    p.add_run("先对每个连续观测段在车体系下减去该段中位数，得到局部磁异常；再根据小车正向/反向转换到统一轨道坐标系。")
    p = doc.add_paragraph()
    p.add_run("重要发现：").bold = True
    p.add_run("4.14 的 X anomaly 量级约几十 nT，而 5.13 的 X anomaly 可达几千到上万 nT；同一 nT 纵轴对比会把 4.14 压成近似横线，因此报告中使用标准化特征图比较形状，并用分面图检查绝对量级。")

    add_picture(
        doc,
        PROC_DIR / "compare_4_14_5_13_x.png",
        "图 1  标准化 X anomaly 对比。用于比较形状特征，避免被跨日期量级差压扁。",
    )
    add_picture(
        doc,
        PROC_DIR / "compare_4_14_5_13_x_absolute_panels.png",
        "图 2  X anomaly 绝对量级检查。上下分面分别使用独立纵轴。",
    )

    doc.add_heading("4. 跨日期磁图对比", level=1)
    add_picture(
        doc,
        PROC_DIR / "compare_4_14_5_13_total.png",
        "图 3  Total 总场磁图对比。Total 对小车方向符号不敏感，是当前最稳的主匹配特征。",
    )
    add_picture(
        doc,
        PROC_DIR / "compare_4_14_5_13_y.png",
        "图 4  标准化 Y anomaly 对比。Y anomaly 在当前指标中跨日期相似性最高。",
    )
    add_picture(
        doc,
        PROC_DIR / "compare_4_14_5_13_z.png",
        "图 5  标准化 Z anomaly 对比。Z anomaly 目前稳定性较差，暂不建议作为主匹配特征。",
    )

    doc.add_heading("5. 相似度指标", level=1)
    doc.add_paragraph(
        "本阶段主要使用以下指标评价两类问题：一是 4.14 与 5.13 两张磁图在同一距离点上的特征相似性；"
        "二是用一段待定位磁曲线在参考磁图上滑动匹配时的定位效果。"
    )
    doc.add_heading("5.1 指标定义", level=2)
    add_bullets(
        doc,
        [
            "Pearson 相关系数 r：衡量两条磁特征曲线形状是否一致，取值范围为 [-1, 1]，越接近 1 表示同向线性相似性越强。公式为 r = Σ[(x_i - x̄)(y_i - ȳ)] / sqrt(Σ(x_i - x̄)^2 Σ(y_i - ȳ)^2)。",
            "梯度相关系数：先对曲线沿距离方向求一阶差分/梯度，再计算 Pearson 相关系数。它更关注局部起伏、峰谷和边缘特征是否一致。",
            "去偏置 RMSE：先去掉两条曲线之间的中位差，再计算均方根误差，用于评价扣除整体基线差后的幅值差异。公式为 RMSE = sqrt(mean((x_i - y_i - median(x-y))^2))。",
            "最佳平移相关：在一定距离范围内平移其中一条曲线，寻找 Pearson 相关系数最大的平移量，用于检查两次建图是否存在整体零点偏差。",
            "NCC 峰值间隔：滑窗匹配时，最佳匹配得分与次优匹配得分的差值。差值越大，说明匹配峰越尖锐，位置歧义越小。",
        ],
    )
    doc.add_heading("5.2 跨日期相似度结果", level=2)
    sim = pd.read_csv(PROC_DIR / "similarity_4_14_vs_5_13.csv")
    add_dataframe_table(
        doc,
        sim,
        [
            "component",
            "overlap_m",
            "pearson_r_same_distance",
            "derivative_r_same_distance",
            "bias_removed_rmse_nT",
            "best_lag_m_for_max_corr",
            "best_lag_pearson_r",
        ],
        ["分量", "重叠/m", "同距相关", "梯度相关", "去偏RMSE/nT", "最佳平移/m", "平移后相关"],
    )
    p = doc.add_paragraph()
    p.add_run("解读：").bold = True
    p.add_run("Y anomaly 的同距相关约 0.73，X anomaly 约 0.59，Total 约 0.54；Z anomaly 暂不稳定。")

    doc.add_heading("6. 初步匹配验证", level=1)
    doc.add_paragraph(
        "匹配验证采用滑动窗口归一化互相关（Sliding-window NCC）方法。参考磁图使用 4.14 融合 total 磁图，"
        "待定位曲线使用 5.13 各连续观测段的 total 曲线。"
    )
    doc.add_heading("6.1 匹配方法", level=2)
    add_numbered(
        doc,
        [
            "从 5.13 某一连续观测段中截取长度为 L 的待匹配窗口，本文测试 L = 20 m、50 m、100 m、150 m。",
            "将该窗口和 4.14 参考磁图中的每一个候选窗口分别做 z-score 标准化，即 x'_i = (x_i - mean(x)) / std(x)。这一步消除常值磁场基线差。",
            "计算待匹配窗口与每个候选窗口的 NCC 得分。由于两条曲线已标准化，NCC 等价于 Pearson 相关系数：NCC(k) = (1/N) Σ q'_i m'_{i+k}。",
            "选择 NCC 得分最大的候选起点作为预测位置，并与该窗口由 SPAN 给出的真实起点比较，得到定位误差。",
            "统计不同窗口长度下的中位绝对误差、P75/P90 误差、NCC 峰值和峰值间隔，用于判断匹配稳定性和位置歧义。",
        ],
    )
    p = doc.add_paragraph()
    p.add_run("说明：").bold = True
    p.add_run(
        "当前验证是在距离域完成的，即假定待匹配曲线已经可以按距离重采样。实际在线定位时，还需要轮速/里程计或 DTW 等方法处理速度变化导致的时间-距离伸缩。"
    )
    doc.add_heading("6.2 匹配验证结果", level=2)
    match = pd.read_csv(PROC_DIR / "matching_validation_summary_total.csv")
    add_dataframe_table(
        doc,
        match,
        [
            "window_m",
            "query_count",
            "median_error_bias_m",
            "median_abs_error_m",
            "p75_abs_error_m",
            "median_best_score",
            "median_score_gap",
        ],
        ["窗口/m", "样本数", "中位偏差/m", "中位绝对误差/m", "P75误差/m", "NCC峰值", "峰值间隔"],
    )
    add_picture(
        doc,
        PROC_DIR / "matching_error_vs_window_total.png",
        "图 6  Total 滑窗 NCC 匹配误差随窗口长度变化。长窗口能降低歧义，但误差仍较大。",
    )
    add_picture(
        doc,
        PROC_DIR / "matching_example_total.png",
        "图 7  Total 匹配示例。局部窗口可匹配到接近位置，但全局仍存在相似片段歧义。",
    )

    doc.add_heading("7. 阶段性结论与下一步建议", level=1)
    add_bullets(
        doc,
        [
            "当前数据可以观察到沿轨道方向的可重复磁特征，说明开展地磁匹配验证是可行的。",
            "仅使用 total 做朴素滑窗 NCC 时，150 m 窗口中位误差约 55 m，说明还有较强位置歧义。",
            "三轴分量不能直接用 raw 值融合；需要考虑小车方向、车体系偏置、姿态角和硬/软铁标定。",
            "短期建议：以 total + total 梯度 + Y anomaly 作为组合特征，加入运动连续性约束，避免匹配位置跳变。",
            "中期建议：做磁强计椭球标定、车体磁补偿和姿态变换，再评估 xyz 是否能显著提升定位精度。",
        ],
    )

    doc.add_heading("附：主要输出文件", level=1)
    add_key_value_table(
        doc,
        [
            ("融合磁图", "magmap_4_14_fused_0p5m.csv；magmap_5_13_fused_0p5m.csv"),
            ("跨日期相似度", "similarity_4_14_vs_5_13.csv"),
            ("匹配验证", "matching_validation_summary_total.csv；matching_validation_5_13_on_4_14_total.csv"),
            ("关键图片", "compare_4_14_5_13_*.png；matching_*.png"),
        ],
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
