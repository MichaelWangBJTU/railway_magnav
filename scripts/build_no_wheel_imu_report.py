from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path.home() / "Desktop" / "\u78c1\u5bfc\u822a" / "\u6570\u636e" / "codex_railway_magnav"
EXP_ROOT = PROJECT_ROOT / "no_wheel_imu"
OUT_DOCX = EXP_ROOT / "reports" / "\u65e0\u8f6e\u901f\u8ba1IMU\u8f85\u52a9\u5730\u78c1\u5b9a\u4f4d\u5b9e\u9a8c\u62a5\u544a.docx"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, df: pd.DataFrame, title: str, cols: list[str], widths: list[float] | None = None) -> None:
    doc.add_paragraph(title, style="Caption")
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        set_cell_text(hdr[i], c, True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            val = row.get(c, "")
            if isinstance(val, float):
                val = f"{val:.3f}"
            set_cell_text(cells[i], val)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def add_picture_if_exists(doc: Document, path: Path, caption: str) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(6.2))
        cap = doc.add_paragraph(caption, style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 78, 121)),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(5)

    styles["Caption"].font.name = "Microsoft YaHei"
    styles["Caption"].font.size = Pt(8.5)
    styles["Caption"].font.color.rgb = RGBColor(90, 90, 90)


def load_outputs() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    outputs = EXP_ROOT / "outputs"
    summary = pd.read_csv(outputs / "no_wheel_imu_matching_summary.csv")
    results = pd.read_csv(outputs / "no_wheel_imu_matching_results.csv")
    quality = pd.read_csv(outputs / "inspvax_relative_distance_quality.csv")
    confidence = pd.read_csv(outputs / "no_wheel_imu_confidence_summary.csv")
    return summary, results, quality, confidence


def fmt_num(x, digits: int = 2) -> str:
    try:
        return f"{float(x):.{digits}f}"
    except Exception:
        return str(x)


def build_report() -> Path:
    summary, results, quality, confidence = load_outputs()
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("\u65e0\u8f6e\u901f\u8ba1\u6761\u4ef6\u4e0b\u7684 IMU \u8f85\u52a9\u94c1\u8def\u5730\u78c1\u5b9a\u4f4d\u8bd5\u9a8c\u62a5\u544a")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 78, 121)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("4.14 \u78c1\u56fe\u4f5c\u4e3a\u53c2\u8003\uff0c5.13 \u91c7\u96c6\u6570\u636e\u4f5c\u4e3a\u67e5\u8be2\uff1b\u4ec5\u4f7f\u7528\u78c1\u5f3a\u8ba1\u4e0e SPAN/INS \u6570\u636e\uff0c\u4e0d\u4f7f\u7528\u8f6e\u901f\u8ba1")

    doc.add_heading("\u7ed3\u8bba\u6458\u8981", level=1)
    best = summary.sort_values("median_abs_error_m").iloc[0]
    doc.add_paragraph(
        "\u672c\u8f6e\u8bd5\u9a8c\u7684\u6838\u5fc3\u7ed3\u8bba\u662f\uff1a\u73b0\u6709 INSPVAX \u901f\u5ea6\u4fe1\u606f\u4e0d\u80fd\u76f4\u63a5\u66ff\u4ee3\u8f6e\u901f\u8ba1\u3002"
        "\u5728\u5254\u9664\u8fc7\u77ed\u8fc7\u6e21\u6bb5\u3001\u6539\u7528\u771f\u5b9e Pearson \u76f8\u5173\u7cfb\u6570\u540e\uff0c\u6700\u597d\u7684\u5168\u8986\u76d6\u65b9\u6cd5\u662f "
        f"{best['method']}\uff0c\u4e2d\u4f4d\u7edd\u5bf9\u8bef\u5dee {fmt_num(best['median_abs_error_m'])} m\uff0c"
        f"RMSE {fmt_num(best['rmse_m'])} m\u3002\u8fd9\u8ddd\u79bb\u8bba\u6587\u4e2d\u5e26\u91cc\u7a0b/\u8f6e\u901f\u7684\u94c1\u8def Graph SLAM \u7c73\u7ea7\u6548\u679c\u8fd8\u5f88\u8fdc\u3002"
    )
    doc.add_paragraph(
        "\u4f46\u5b9e\u9a8c\u4e5f\u5f97\u5230\u4e86\u6709\u4ef7\u503c\u7684\u65b9\u5411\uff1a\u65e0\u8f6e\u901f\u8ba1\u6761\u4ef6\u4e0b\uff0c\u4e0d\u5e94\u628a INSPVAX \u6c34\u5e73\u901f\u5ea6\u79ef\u5206\u5f53\u4f5c\u5f3a\u91cc\u7a0b\uff0c"
        "\u800c\u5e94\u628a IMU/INS \u7528\u4f5c\u65b9\u5411\u3001\u59ff\u6001\u3001\u5f31\u91cc\u7a0b\u6216\u5c3a\u5ea6\u5148\u9a8c\uff0c\u518d\u901a\u8fc7\u78c1\u7279\u5f81\u5339\u914d\u4e0e\u552f\u4e00\u6027\u95e8\u9650\u505a\u7edd\u5bf9\u4f4d\u7f6e\u6821\u6b63\u3002"
    )

    doc.add_heading("\u4e0e\u6587\u732e\u65b9\u6cd5\u7684\u5173\u7cfb", level=1)
    doc.add_paragraph(
        "Siebler \u7b49\u4eba\u5728 FUSION 2024 \u63d0\u51fa\u94c1\u8def\u573a\u666f\u7684\u5730\u78c1 Graph SLAM\uff1a\u6bcf\u4e2a\u56fe\u8282\u70b9\u4fdd\u5b58\u5c40\u90e8\u78c1\u56fe\uff0c"
        "\u7528\u7ea6\u6570\u5341\u7c73\u7684\u78c1\u7b7e\u540d\u505a loop closure\uff0c\u518d\u628a loop closure \u4f5c\u4e3a pose graph \u7ea6\u675f\u3002"
        "\u8fd9\u7c7b\u65b9\u6cd5\u7684\u524d\u63d0\u662f\u5b58\u5728\u53ef\u9760\u91cc\u7a0b\u8ba1/\u8f6e\u901f\u8ba1\u3002\u4f60\u7684\u5c0f\u8f66\u6682\u65f6\u6ca1\u6709\u8f6e\u901f\u8ba1\uff0c\u6240\u4ee5\u4e0d\u80fd\u76f4\u63a5\u590d\u5236\u8be5 SOTA \u7684\u8f93\u5165\u6761\u4ef6\u3002"
    )
    doc.add_paragraph(
        "ION GNSS+ 2017 \u7684 INS+\u78c1\u7b7e\u540d\u65b9\u6cd5\u66f4\u63a5\u8fd1\u672c\u95ee\u9898\uff1a\u7528\u78c1\u56fe\u5339\u914d\u7ed9 INS \u63d0\u4f9b\u4f4d\u7f6e\u66f4\u65b0\uff0c\u5728 ESKF \u4e2d\u7ea6\u675f INS \u6f02\u79fb\u3002"
        "\u8fd9\u63d0\u793a\u6211\u4eec\u540e\u7eed\u5e94\u628a\u539f\u59cb IMU \u9884\u79ef\u5206\u3001\u94c1\u8def\u4e00\u7ef4\u8fd0\u52a8\u7ea6\u675f\u3001\u78c1\u5339\u914d\u56e0\u5b50\u653e\u5230\u4e00\u4e2a\u56e0\u5b50\u56fe/\u6ee4\u6ce2\u5668\u91cc\u3002"
    )
    doc.add_paragraph(
        "WM-GFM \u5219\u662f\u5f31\u91cc\u7a0b\u8f85\u52a9\u7684\u5730\u78c1\u7279\u5f81\u5339\u914d\u601d\u8def\uff1a\u5b83\u4e0d\u5b8c\u5168\u4f9d\u8d56\u9ad8\u7cbe\u91cc\u7a0b\uff0c\u800c\u662f\u7528\u5f31\u91cc\u7a0b\u7ea6\u675f\u68af\u5ea6\u7279\u5f81\u548c\u5339\u914d\u8def\u5f84\u3002"
        "\u672c\u8f6e\u5b9e\u9a8c\u5c1d\u8bd5\u628a INSPVAX \u901f\u5ea6\u79ef\u5206\u4f5c\u4e3a\u5f31\u91cc\u7a0b\uff0c\u4f46\u5b9e\u9a8c\u663e\u793a\u5176\u5c3a\u5ea6\u5728\u4e0d\u540c\u7247\u6bb5\u4e0d\u7a33\u5b9a\u3002"
    )

    doc.add_heading("\u6570\u636e\u4e0e\u9884\u5904\u7406", level=1)
    doc.add_paragraph(
        "\u53c2\u8003\u56fe\u4f7f\u7528 4.14 \u878d\u5408\u78c1\u56fe\uff0c\u67e5\u8be2\u65e5\u4f7f\u7528 5.13 \u539f\u59cb\u65f6\u95f4\u987a\u5e8f\u7684\u78c1\u4f20\u611f\u5668\u6837\u672c\u3002"
        "\u672c\u6b21\u4e0d\u518d\u4ece 0.5 m \u7f51\u683c\u78c1\u56fe\u53cd\u63a8\u65f6\u95f4\uff0c\u56e0\u4e3a\u8fd4\u7a0b\u7247\u6bb5\u4e2d\u7f51\u683c\u65f6\u95f4\u53ef\u80fd\u5c40\u90e8\u5012\u5e8f\uff0c\u4f1a\u7834\u574f\u901f\u5ea6\u79ef\u5206\u3002"
        "\u5b9e\u9a8c\u4e2d\u5254\u9664\u4e86\u8fc7\u77ed\u7684\u8fc7\u6e21/\u8c03\u5934\u7247\u6bb5\uff0c\u6700\u7ec8\u7528 5 \u4e2a\u6709\u6548\u7247\u6bb5\u8bc4\u4ef7\u3002"
    )
    doc.add_paragraph(
        "INSPVAX \u65f6\u95f4\u5904\u7406\uff1aNovAtel \u65e5\u5fd7\u5934\u4e2d\u7684 GPS week/seconds-of-week \u6309 GPS \u65f6\u95f4\u89e3\u6790\uff0c"
        "\u5148\u51cf 18 s \u95f0\u79d2\u5f97 UTC\uff0c\u518d\u52a0 8 h \u5f97\u5317\u4eac\u65f6\u95f4\uff0c\u4e0e\u78c1\u5f3a\u8ba1\u7cfb\u7edf\u65f6\u95f4\u5bf9\u9f50\u3002"
    )

    q = quality.copy()
    q = q[["segment_label", "direction", "true_length_m", "imu_length_m", "rel_scale", "rel_rmse_m", "mean_speed_mps", "median_vel_std_mps"]]
    q.columns = ["\u7247\u6bb5", "\u65b9\u5411", "\u771f\u503c\u957f\u5ea6/m", "INSPVAX\u79ef\u5206\u957f\u5ea6/m", "\u7ebf\u6027\u5c3a\u5ea6", "\u76f8\u5bf9\u91cc\u7a0bRMSE/m", "\u5e73\u5747\u901f\u5ea6/m/s", "\u901f\u5ea6\u6807\u51c6\u5dee\u4e2d\u4f4d/m/s"]
    add_table(doc, q, "\u8868 1  INSPVAX \u76f8\u5bf9\u91cc\u7a0b\u8d28\u91cf\u8bca\u65ad", list(q.columns))

    doc.add_heading("\u65b9\u6cd5\u8bbe\u8ba1", level=1)
    doc.add_heading("1. \u78c1\u7279\u5f81", level=2)
    doc.add_paragraph(
        "\u5bf9\u603b\u573a total \u548c\u8f68\u9053\u5750\u6807\u7cfb Y \u5411\u5f02\u5e38\u91cf\u6784\u9020\u9ad8\u901a\u7279\u5f81\uff1a"
        "m_HP(s)=m(s)-median(m(u), |u-s|<=W/2), \u672c\u6b21 W=30 m\u3002\u603b\u573a\u8fd8\u8ba1\u7b97\u68af\u5ea6 g(s)=dm_HP(s)/ds\u3002"
        "\u5339\u914d\u65f6\u4f7f\u7528\u603b\u573a\u9ad8\u901a\u3001\u603b\u573a\u68af\u5ea6\u3001Y \u5411\u9ad8\u901a\u4e09\u7c7b\u7279\u5f81\u7684\u5e73\u5747\u76f8\u5173\u5ea6\u3002"
    )
    doc.add_paragraph(
        "\u76f8\u4f3c\u5ea6\u91c7\u7528 Pearson \u76f8\u5173\u7cfb\u6570\uff1a"
        "\u03c1(a,b)=\u03a3_i(a_i-\u0061\u0304)(b_i-\u0062\u0304)/sqrt(\u03a3_i(a_i-\u0061\u0304)^2\u03a3_i(b_i-\u0062\u0304)^2)\u3002"
        "\u4e3a\u4e86\u964d\u4f4e\u5f02\u5e38\u503c\u5f71\u54cd\uff0c\u5728\u8ba1\u7b97\u524d\u5148\u7528 median/MAD \u505a\u7a33\u5065\u6807\u51c6\u5316\u5e76\u9650\u5e45\u3002"
    )

    doc.add_heading("2. \u65e0\u8f6e\u901f\u8ba1\u65f6\u95f4\u5c3a\u5ea6\u641c\u7d22", level=2)
    doc.add_paragraph(
        "\u8fd9\u662f\u4e0d\u4f7f\u7528\u8f6e\u901f\u4e5f\u4e0d\u76f4\u63a5\u4f7f\u7528 INSPVAX \u901f\u5ea6\u7684 baseline\u3002"
        "\u5bf9\u4e00\u6bb5\u67e5\u8be2\u78c1\u5e8f\u5217\uff0c\u5047\u8bbe\u76f8\u5bf9\u91cc\u7a0b\u4e0e\u65f6\u95f4\u6210\u6bd4\u4f8b\uff1ar_i=L(t_i-t_0)/(t_N-t_0)\u3002"
        "\u7b97\u6cd5\u641c\u7d22\u603b\u957f\u5ea6 L \u548c\u5730\u56fe\u8d77\u70b9 p0\uff0c\u5728\u524d\u8fdb/\u8fd4\u56de\u65b9\u5411 d \u4e0b\u4ee4 p_i=p0+d*r_i\uff0c\u53d6\u76f8\u5173\u5ea6\u6700\u5927\u7684 p0\u3002"
    )

    doc.add_heading("3. INSPVAX \u901f\u5ea6\u79ef\u5206\u4e0e\u5c3a\u5ea6\u641c\u7d22", level=2)
    doc.add_paragraph(
        "\u76f4\u63a5\u901f\u5ea6\u79ef\u5206\u65b9\u6cd5\u4ee4 r_i=\u222b sqrt(v_N^2+v_E^2)dt\uff0c\u5176\u4e2d v_N,v_E \u6765\u81ea INSPVAX\u3002"
        "\u8003\u8651\u5230 SPAN/INS \u901f\u5ea6\u53ef\u80fd\u6709\u5c3a\u5ea6\u504f\u5dee\uff0c\u53c8\u6d4b\u8bd5\u4e86 r_i=\u03b1\u222b sqrt(v_N^2+v_E^2)dt\uff0c\u5bf9\u03b1 \u548c p0 \u4e00\u8d77\u641c\u7d22\u3002"
    )

    doc.add_heading("4. \u552f\u4e00\u6027\u95e8\u9650\u4e0e\u7a97\u53e3\u5339\u914d", level=2)
    doc.add_paragraph(
        "\u4e3a\u4e86\u51cf\u5c11\u9519\u8bef\u5339\u914d\uff0c\u8ba1\u7b97\u6700\u4f73\u76f8\u5173\u5cf0\u4e0e\u8ddd\u5176\u81f3\u5c11 20 m \u7684\u7b2c\u4e8c\u9ad8\u5cf0\u4e4b\u5dee\uff1a"
        "\u0394\u03c1=\u03c1_best-max(\u03c1(p), |p-p_best|>=20m)\u3002"
        "\u5f53 \u0394\u03c1 \u8fc7\u5c0f\u65f6\uff0c\u8bf4\u660e\u5730\u56fe\u4e0a\u6709\u591a\u4e2a\u76f8\u4f3c\u5019\u9009\uff0c\u5e94\u62d2\u7edd\u5b9a\u4f4d\u3002\u672c\u6b21\u4e5f\u5c1d\u8bd5\u4e86 120 m \u5b50\u7a97\u53e3\u5339\u914d\uff0c\u5bfb\u627e\u66f4\u5177\u8fa8\u8bc6\u5ea6\u7684\u5c40\u90e8\u78c1\u7b7e\u540d\u3002"
    )

    doc.add_heading("\u5b9e\u9a8c\u7ed3\u679c", level=1)
    s = summary.copy()
    s = s[["method", "segment_count", "median_abs_error_m", "mean_abs_error_m", "rmse_m", "p90_abs_error_m", "median_score", "median_margin", "accepted_0p20"]]
    s.columns = ["\u65b9\u6cd5", "\u7247\u6bb5\u6570", "\u4e2d\u4f4d\u7edd\u5bf9\u8bef\u5dee/m", "\u5e73\u5747\u7edd\u5bf9\u8bef\u5dee/m", "RMSE/m", "P90\u7edd\u5bf9\u8bef\u5dee/m", "\u4e2d\u4f4d\u76f8\u5173", "\u4e2d\u4f4d\u95e8\u9650\u5dee", "\u0394\u03c1>=0.20\u6570"]
    add_table(doc, s, "\u8868 2  \u65e0\u8f6e\u901f\u8ba1/IMU \u65b9\u6cd5\u5b9a\u4f4d\u8bef\u5dee\u6c47\u603b", list(s.columns))

    selected_conf = confidence[
        (confidence["margin_threshold"].isin([0.0, 0.05, 0.08, 0.1]))
        & (confidence["accepted_count"] > 0)
    ].copy()
    selected_conf = selected_conf[["method", "margin_threshold", "accepted_count", "total_count", "coverage_pct", "median_abs_error_m", "mean_abs_error_m"]]
    selected_conf.columns = ["\u65b9\u6cd5", "\u95e8\u9650", "\u63a5\u53d7\u6570", "\u603b\u6570", "\u8986\u76d6\u7387/%", "\u63a5\u53d7\u540e\u4e2d\u4f4d\u8bef\u5dee/m", "\u63a5\u53d7\u540e\u5e73\u5747\u8bef\u5dee/m"]
    add_table(doc, selected_conf, "\u8868 3  \u552f\u4e00\u6027\u95e8\u9650\u7684\u8986\u76d6\u7387-\u7cbe\u5ea6\u6743\u8861", list(selected_conf.columns))

    fig_dir = EXP_ROOT / "figures"
    add_picture_if_exists(doc, fig_dir / "no_wheel_imu_method_summary.png", "\u56fe 1  \u5404\u65b9\u6cd5\u4e2d\u4f4d\u7edd\u5bf9\u8bef\u5dee\u5bf9\u6bd4")
    add_picture_if_exists(doc, fig_dir / "inspvax_relative_distance_quality.png", "\u56fe 2  INSPVAX \u76f8\u5bf9\u91cc\u7a0b\u8d28\u91cf\u8bca\u65ad")
    add_picture_if_exists(doc, fig_dir / "no_wheel_imu_confidence_tradeoff.png", "\u56fe 3  \u552f\u4e00\u6027\u95e8\u9650\u4e0b\u7684\u8986\u76d6\u7387\u4e0e\u8bef\u5dee\u53d8\u5316")

    doc.add_heading("\u95ee\u9898\u5206\u6790", level=1)
    doc.add_paragraph(
        "1. INSPVAX \u901f\u5ea6\u5c3a\u5ea6\u4e0d\u7a33\u5b9a\u3002\u4e0d\u540c\u7247\u6bb5\u7684\u7ebf\u6027\u5c3a\u5ea6\u4ece\u7ea6 0.45 \u5230 2.18 \u90fd\u51fa\u73b0\uff0c"
        "\u8bf4\u660e\u5b83\u4e0d\u80fd\u4f5c\u4e3a\u7a33\u5b9a\u91cc\u7a0b\u8ba1\u3002\u8fd9\u53ef\u80fd\u4e0e\u5c0f\u8f66\u4f4e\u901f\u3001\u505c\u8d70\u3001INS/GNSS \u878d\u5408\u72b6\u6001\u548c\u8f68\u9053\u65b9\u5411\u6295\u5f71\u672a\u7cbe\u7ec6\u5904\u7406\u6709\u5173\u3002"
    )
    doc.add_paragraph(
        "2. \u78c1\u7279\u5f81\u5728 500-600 m \u8303\u56f4\u5185\u5b58\u5728\u81ea\u76f8\u4f3c\u7247\u6bb5\u3002\u5373\u4f7f\u6700\u4f73\u76f8\u5173\u770b\u8d77\u6765\u4e0d\u4f4e\uff0c\u4e5f\u53ef\u80fd\u5339\u5230\u9519\u8bef\u4f4d\u7f6e\u3002"
        "\u56e0\u6b64\u4e0d\u80fd\u53ea\u770b\u6700\u5927\u76f8\u5173\u503c\uff0c\u5fc5\u987b\u52a0\u552f\u4e00\u6027\u3001\u8fd0\u52a8\u5148\u9a8c\u548c\u9c81\u68d2\u56e0\u5b50\u3002"
    )
    doc.add_paragraph(
        "3. \u76f4\u63a5\u590d\u73b0\u8f6e\u901f\u8ba1\u7248 SOTA \u4e0d\u516c\u5e73\u3002FUSION 2024 \u94c1\u8def Graph SLAM \u7684\u6838\u5fc3\u662f odometer+magnetometer\uff0c"
        "\u800c\u672c\u6570\u636e\u6ca1\u6709\u8f6e\u901f\u8ba1\u3002\u5982\u679c\u5f3a\u884c\u7528 INSPVAX \u901f\u5ea6\u4ee3\u66ff\uff0c\u5c31\u4f1a\u628a INS/GNSS \u878d\u5408\u8bef\u5dee\u5f15\u5165\u91cc\u7a0b\u8fb9\uff0c\u5bfc\u81f4\u56fe\u4f18\u5316\u524d\u7aef\u7684 loop closure \u5019\u9009\u672c\u8eab\u5c31\u4e0d\u53ef\u9760\u3002"
    )

    doc.add_heading("\u5efa\u8bae\u7684\u4e0b\u4e00\u6b65\u521b\u65b0\u65b9\u5411", level=1)
    doc.add_paragraph(
        "\u5efa\u8bae\u5c06\u540e\u7eed\u7814\u7a76\u5b9a\u4e49\u4e3a\uff1a\u201c\u65e0\u8f6e\u901f\u8ba1\u94c1\u8def\u573a\u666f\u4e0b\u7684 IMU \u5f31\u7ea6\u675f\u78c1\u7b7e\u540d\u56e0\u5b50\u56fe\u5b9a\u4f4d\u201d\u3002"
        "\u72b6\u6001\u91cf\u53ef\u8bbe\u4e3a\u6cbf\u8f68\u4f4d\u7f6e p\u3001\u901f\u5ea6 v\u3001IMU bias\u3001\u91cc\u7a0b\u5c3a\u5ea6\u03b1\u3002"
        "\u8fd0\u52a8\u8fb9\u6765\u81ea IMU \u9884\u79ef\u5206\u548c\u94c1\u8def\u4e00\u7ef4\u8fd0\u52a8\u7ea6\u675f\uff0c\u78c1\u5339\u914d\u8fb9\u6765\u81ea\u9ad8\u552f\u4e00\u6027\u7684\u78c1\u7b7e\u540d\u3002"
    )
    doc.add_paragraph(
        "\u4e0b\u4e00\u8f6e\u5efa\u8bae\u8f6c\u6362 RAWIMUSX \u6216\u76f8\u5e94\u539f\u59cb IMU \u65e5\u5fd7\uff0c\u4e0d\u53ea\u7528 INSPVAX \u7ed3\u679c\u901f\u5ea6\u3002"
        "\u5177\u4f53\u53ef\u505a\uff1a\u6cbf\u8f68\u65b9\u5411\u6295\u5f71\u52a0\u901f\u5ea6\uff0c\u7528\u505c\u8f66\u7247\u6bb5\u505a ZUPT/\u96f6\u901f\u66f4\u65b0\uff0c\u5728\u56e0\u5b50\u56fe\u4e2d\u540c\u65f6\u4f30\u8ba1 bias \u548c\u5c3a\u5ea6\uff0c"
        "\u5e76\u5c06\u78c1\u5339\u914d\u56e0\u5b50\u8bbe\u6210 switchable/robust factor\uff0c\u4f7f\u9519\u8bef loop closure \u4e0d\u4f1a\u62d6\u57ae\u8f68\u8ff9\u3002"
    )
    doc.add_paragraph(
        "\u8fd9\u4e2a\u65b9\u5411\u7684\u4ef7\u503c\u5728\u4e8e\uff1a\u5b83\u4e0d\u662f\u7b80\u5355\u52a0\u4e00\u4e2a\u65b9\u5411\u4fe1\u606f\uff0c\u800c\u662f\u6b63\u9762\u89e3\u51b3\u201c\u65e0\u8f6e\u901f\u8ba1\u65f6\u7a7a\u95f4\u5c3a\u5ea6\u4e0d\u7a33\u5b9a\u201d\u95ee\u9898\u3002"
        "\u5982\u679c\u80fd\u5728\u5f53\u524d 4.14/5.13 \u6570\u636e\u4e0a\u5c06\u63a5\u53d7\u7247\u6bb5\u8bef\u5dee\u538b\u5230 10-30 m \u5185\uff0c\u518d\u901a\u8fc7\u591a\u8d9f\u56e0\u5b50\u56fe\u7d2f\u79ef\u4f18\u5316\uff0c\u5c31\u6709\u53ef\u80fd\u5f62\u6210\u4e00\u4e2a\u9488\u5bf9\u65e0\u8f6e\u901f\u8ba1\u94c1\u8def\u5e73\u53f0\u7684\u521b\u65b0\u70b9\u3002"
    )

    doc.add_heading("\u53c2\u8003\u6587\u732e", level=1)
    refs = [
        "Siebler, B., Lehner, A., Sand, S., Hanebeck, U. D. Magnetic Field Mapping of Railway Lines with Graph SLAM. FUSION 2024. DOI: 10.23919/FUSION59988.2024.10706392. https://isas.iar.kit.edu/pdf/FUSION24_Siebler.pdf",
        "Siebler, B., Heirich, O., Sand, S. Bounding INS Positioning Errors with Magnetic-Field-Signatures in Railway Environments. ION GNSS+ 2017. DOI: 10.33012/2017.15311. https://www.ion.org/publications/abstract.cfm?articleID=15311",
        "You, Y. et al. WM-GFM: a novel geomagnetic feature matching positioning method with weak mileage aid. Measurement, 257, 118850, 2026. DOI: 10.1016/j.measurement.2025.118850. https://www.sciencedirect.com/science/article/pii/S0263224125022092",
        "Kok, M., Solin, A. Online One-Dimensional Magnetic Field SLAM with Loop-Closure Detection. arXiv:2409.01091, 2024. https://arxiv.org/abs/2409.01091",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build_report())
