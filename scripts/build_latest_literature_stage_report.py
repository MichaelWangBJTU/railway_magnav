from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\m1352\Documents\railway_magnav")
OUT = ROOT / "reports"
DOCX = OUT / "无轮速计铁路地磁定位阶段性研究报告_20260609.docx"


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size_pt: float, color: str = "000000", bold: bool = False) -> None:
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size_pt)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "000000") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    set_run_font(run, 8.5, bold, color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, 10.5)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, 10.5)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, 10.5)


def add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("1F4D78")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Cm(widths_cm[i])
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True, color="1F4D78")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Cm(widths_cm[i])
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_picture_if_exists(doc: Document, path: Path, caption: str, width_in: float = 6.3) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_run_font(r, 8.5, False, "555555")


def best_rows_from_csv(path: Path, keep: list[str], limit: int = 8) -> list[list[str]]:
    if not path.exists():
        return []
    df = pd.read_csv(path)
    rows = []
    for _, row in df.head(limit).iterrows():
        rows.append([str(row.get(col, "")) for col in keep])
    return rows


def build() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    set_style_font(styles["Normal"], 10.5)
    set_style_font(styles["Heading 1"], 16, "2E74B5", True)
    set_style_font(styles["Heading 2"], 13, "2E74B5", True)
    set_style_font(styles["Heading 3"], 11.5, "1F4D78", True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("无轮速计铁路地磁定位阶段性研究报告")
    set_run_font(r, 20, True, "0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("数据：4.14 建图 / 5.13 跨日验证；日期：2026-06-09")
    set_run_font(r, 10.5, False, "555555")

    doc.add_heading("1. 一句话结论", level=1)
    add_paragraph(
        doc,
        "最新文献表明，铁路地磁 SOTA 往往依赖里程计或高质量空间磁阵列；在我们无轮速计、短线路、跨日轴不一致的数据条件下，单纯 MSD/ICCP 或短序列全局匹配仍会被重复磁签名骗到。当前最有价值的路线是：总场高通基线 + 跨日轴/标定自适应 + HMM/Viterbi 或粒子滤波运动约束 + 置信拒绝。"
    )

    doc.add_heading("2. 传感器条件与公平 baseline", level=1)
    add_bullets(
        doc,
        [
            "当前没有轮速计或可靠里程计；可以使用磁强计、SPAN/GNSS 真值做离线评估，也可以尝试使用 INSPVAX/BESTVEL 速度作为弱先验。",
            "FUSION 2024 Graph SLAM 是强 SOTA 参考，但其核心边约束依赖 odometer，因此不能直接作为我们当前条件下的公平 baseline。",
            "更公平的 baseline 是 FUSION 2018 风格的铁路磁图粒子滤波，以及其确定性 HMM/Viterbi 近似版本。",
        ],
    )

    doc.add_heading("3. 近期文献定位", level=1)
    lit_rows = [
        ["2018", "Train Localization with Particle Filter and Magnetic Field Measurements", "磁图 + 磁强计 + 运动模型", "无轮速计最接近 baseline；文献 RMSE 3.84 m"],
        ["2024", "Magnetic Field Mapping of Railway Lines with Graph SLAM", "磁强计 + odometer + pose graph", "强 SOTA 参考；loop closure RMSE 约 0.45 m，但依赖里程计"],
        ["2025", "Real-time rail vehicle localisation using spatially resolved magnetic field measurements", "空间磁序列 + PF + top-k 初始化", "支持短序列候选生成；其 top-3 初始化很强"],
        ["2025", "Snapshot Estimator with Uncalibrated Magnetometers", "未标定磁强计 + SLAC/snapshot", "支持局部标定/轴自适应思想"],
        ["2026", "Weak-mileage geomagnetic feature matching", "弱里程 + 地磁特征匹配", "支持使用 INSPVAX 速度作为弱里程；需全文后才能复现"],
    ]
    add_table(doc, ["年份", "文献", "条件", "对本项目意义"], lit_rows, [1.2, 5.1, 4.0, 5.4])

    doc.add_heading("4. 已验证的数据问题", level=1)
    add_paragraph(doc, "跨日向量轴不能直接比较。前向片段中，经验关系为：")
    add_formula(doc, "X_4.14 ≈ Z_5.13,   Y_4.14 ≈ Y_5.13,   Z_4.14 ≈ X_5.13")
    add_paragraph(doc, "局部去偏置后，X/Y 曲线可以达到 nT 级相似；但全局搜索仍会出现远距离重复峰。因此师兄局部图看起来很好，并不等价于全局定位已经解决。")
    add_picture_if_exists(
        doc,
        ROOT / "msd_iccp_check" / "cross_day_axis_remapped_debiased_true_4_14_seg5_01_vs_5_13_seg9_02.png",
        "图 1 轴重映射后局部跨日曲线可高度相似，但该结论只说明局部可复现，不保证全局唯一。",
    )

    doc.add_heading("5. 主要方法与公式", level=1)
    doc.add_heading("5.1 短序列匹配指标", level=2)
    add_paragraph(doc, "NCC 用于相似度最大化，MSD 用于差异最小化：")
    add_formula(doc, "NCC(q,m)=Σ((q_i-μ_q)(m_i-μ_m))/(sqrt(Σ(q_i-μ_q)^2) sqrt(Σ(m_i-μ_m)^2))")
    add_formula(doc, "MSD(q,m)=(1/N) Σ(q_i-m_i)^2")
    add_paragraph(doc, "top-k 指标不只看最优候选，而是看正确位置是否保留在前 k 个候选中：")
    add_formula(doc, "success@k,τ = 1{∃ r≤k, |ŝ_r - s_true| ≤ τ}")

    doc.add_heading("5.2 HMM/Viterbi 无轮速计匹配", level=2)
    add_paragraph(doc, "将沿轨位置离散为 0.5 m 状态，使用方向和最大速度限制转移。Viterbi 递推为：")
    add_formula(doc, "D_k(j)=log p(z_k|s_j)+max_i [D_{k-1}(i)+log p(s_j|s_i)]")
    add_paragraph(doc, "观测似然使用 robust z-score 和重尾损失，降低局部异常磁干扰的影响：")
    add_formula(doc, "log p_f ∝ -0.5(ν+1) log(1 + (r_f/σ)^2/ν),   ν=3")

    doc.add_heading("5.3 SLAC-lite 局部轴/标定自适应", level=2)
    add_paragraph(doc, "受 2025 snapshot/SLAC 文献启发，在每个候选位置局部估计查询三轴到参考三轴的仿射变换：")
    add_formula(doc, "min_{A,b} Σ_i || r_i - (A q_i + b) ||^2 + λ||A||_F^2")
    add_paragraph(doc, "该方法能解释跨日轴不一致，但在本数据上也容易把错误候选拟合得很好，因此必须配合运动约束与唯一性判据。")

    doc.add_heading("6. 当前实验结果", level=1)
    hmm_rows = [
        ["Previous RobustTotalHP Viterbi", "62.8", "109.2", "123.1", "早期无轮速计基线"],
        ["AxisCal XY+Total InfoGate Viterbi", "17.6", "91.7", "108.6", "中位数明显改善，但有失败片段"],
        ["AxisCal XY+Total MidGate Viterbi", "17.6", "90.4", "107.3", "当前中位数最好"],
        ["SpeedPrior TotalHP Viterbi", "38.5", "71.5", "78.9", "均值/RMSE 更稳，利用 INSPVAX 速度"],
    ]
    add_table(doc, ["方法", "中位误差/m", "平均误差/m", "RMSE/m", "说明"], hmm_rows, [4.5, 2.0, 2.0, 2.0, 5.0])
    add_picture_if_exists(
        doc,
        ROOT / "axis_calibrated_hmm_speed_prior" / "axis_calibrated_hmm_summary.png",
        "图 2 HMM/Viterbi 方法对比。AxisCal 方法中位数好，SpeedPrior 方法整体稳定性更好。",
    )

    doc.add_heading("7. 最新文献触发的新实验", level=1)
    topk_path = ROOT / "latest_literature_aligned_experiments" / "short_sequence_topk_summary.csv"
    topk_df = pd.read_csv(topk_path)
    view = topk_df[
        (topk_df["window_m"] == 150.0)
        & topk_df["method"].isin(["TotalHP_NCC", "AxisCal_XY_MSD", "SLAC_Affine_OldXYZ_to_RefXYZ"])
    ].copy()
    view = view.sort_values("median_top1_abs_error_m")
    rows = []
    for _, row in view.iterrows():
        rows.append(
            [
                row["method"],
                f'{row["median_top1_abs_error_m"]:.2f}',
                f'{row["top1_within_25m_rate"]:.3f}',
                f'{row["top3_within_25m_rate"]:.3f}',
                f'{row["median_score_gap_10m"]:.3f}',
            ]
        )
    add_table(doc, ["方法", "Top-1中位误差/m", "Top-1@25m", "Top-3@25m", "score gap"], rows, [4.6, 2.6, 2.0, 2.0, 2.0])
    add_picture_if_exists(
        doc,
        ROOT / "latest_literature_aligned_experiments" / "short_sequence_topk_summary.png",
        "图 3 短序列全局初始化结果。正确位置常在候选中，但 top-1 仍受重复特征影响。",
    )

    doc.add_heading("8. 负结果与风险", level=1)
    add_bullets(
        doc,
        [
            "MSD+ICCP 只能作为 baseline。没有可靠初始位置时，ICCP 容易被错误等值/相似特征牵引。",
            "局部轴标定不能单独解决问题；SLAC-lite 的 score gap 很小，说明错误候选也能被局部线性变换拟合。",
            "单看 Viterbi final score margin 选择方法会失败。重复磁签名可以产生高置信假路径。",
            "Graph SLAM 结果不能直接用于宣称超越 SOTA，因为其强约束来自 odometer。",
        ],
    )

    doc.add_heading("9. 下一步创新路线", level=1)
    add_numbered(
        doc,
        [
            "构建 top-k 候选生成器：对 50-150 m 短序列输出候选簇，而不是只输出 top-1。",
            "将候选簇作为 HMM/粒子滤波的观测似然或初始化分布，引入方向、最大速度和 INSPVAX 弱速度约束。",
            "将 SLAC/轴自适应作为候选局部似然的一部分，但加入复杂度惩罚和唯一性门限，避免过拟合错误候选。",
            "设计可靠性层：结合候选峰间距、top-k 分数形状、多方法轨迹分歧、速度/进度一致性，输出定位或拒绝定位。",
            "按 4.14 组内 leave-one-pass-out 和 5.13 跨日验证分别评价，避免只在跨日 5 个片段上过拟合。",
        ],
    )

    doc.add_heading("10. 参考文献", level=1)
    refs = [
        "Siebler et al., Train Localization with Particle Filter and Magnetic Field Measurements, FUSION 2018. https://elib.dlr.de/119898/1/FUSION_2018.pdf",
        "Siebler et al., Robust Particle Filter for Magnetic field-based Train Localization, ION GNSS+ 2022. https://www.ion.org/publications/abstract.cfm?articleID=18536",
        "Siebler et al., Magnetic Field Mapping of Railway Lines with Graph SLAM, FUSION 2024. https://isas.iar.kit.edu/pdf/FUSION24_Siebler.pdf",
        "Dieckow et al., Real-time rail vehicle localisation using spatially resolved magnetic field measurements, arXiv 2025. https://arxiv.org/pdf/2507.19327",
        "Siebler et al., Snapshot Estimator for Magnetic Field-based Train Localization with Uncalibrated Magnetometers, EUSIPCO 2025. https://eusipco2025.org/wp-content/uploads/pdfs/0002142.pdf",
        "A novel geomagnetic feature matching positioning method with weak mileage aid, Measurement 257:118850, 2026. https://www.sciencedirect.com/science/article/abs/pii/S026322412600268X",
    ]
    add_numbered(doc, refs)

    doc.add_section(WD_SECTION_START.CONTINUOUS)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
