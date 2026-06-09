from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


SRC_MD = Path("method_record_no_wheel_magnav.md")
PROJECT_ROOT = Path.home() / "Desktop" / "\u78c1\u5bfc\u822a" / "\u6570\u636e" / "codex_railway_magnav"
OUT_DOCX = PROJECT_ROOT / "no_wheel_sota" / "reports" / "\u65e0\u8f6e\u901f\u8ba1\u94c1\u8def\u5730\u78c1\u5b9a\u4f4d\u65b9\u6cd5\u8bb0\u5f55.docx"


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 78, 121)),
    ]:
        st = doc.styles[name]
        st.font.name = "Microsoft YaHei"
        st.font.size = Pt(size)
        st.font.color.rgb = color


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)


def add_markdown_table(doc: Document, rows: list[str]) -> None:
    parsed = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in rows]
    parsed = [r for r in parsed if not all(set(c) <= {"-", ":"} for c in r)]
    if not parsed:
        return
    table = doc.add_table(rows=1, cols=len(parsed[0]))
    table.style = "Table Grid"
    for i, cell in enumerate(parsed[0]):
        run = table.rows[0].cells[i].paragraphs[0].add_run(cell)
        run.bold = True
        run.font.size = Pt(8.5)
    for row in parsed[1:]:
        cells = table.add_row().cells
        for i, cell in enumerate(row[: len(cells)]):
            cells[i].text = cell
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)


def build() -> Path:
    text = SRC_MD.read_text(encoding="utf-8")
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("\u65e0\u8f6e\u901f\u8ba1\u94c1\u8def\u5730\u78c1\u5b9a\u4f4d\u65b9\u6cd5\u8bb0\u5f55")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)

    code: list[str] | None = None
    table_rows: list[str] = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if code is None:
                code = []
            else:
                add_code_block(doc, code)
                code = None
            continue
        if code is not None:
            code.append(line)
            continue
        if line.startswith("|") and line.endswith("|"):
            table_rows.append(line)
            continue
        if table_rows:
            add_markdown_table(doc, table_rows)
            table_rows = []
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line[:3].replace(".", "").isdigit() and ". " in line[:5]:
            doc.add_paragraph(line.split(". ", 1)[1].strip(), style="List Number")
        else:
            doc.add_paragraph(line)
    if table_rows:
        add_markdown_table(doc, table_rows)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build())
