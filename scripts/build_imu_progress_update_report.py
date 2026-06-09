from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\m1352\Documents\railway_magnav")
OUT = ROOT / "reports" / "无轮速计铁路地磁定位_IMU进度门控集成更新报告_20260609.docx"


def set_font(run, size=10.5, bold=False, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r)
    return p


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True, color="1F4D78")


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r, size=10)


def add_table(doc: Document, df: pd.DataFrame, font_size: float = 8.5) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_font(r, font_size, True, "0B2545")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                text = f"{val:.3f}"
            else:
                text = str(val)
            cells[j].text = text
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_font(r, font_size)
    doc.add_paragraph()


def picture(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_font(r, 9, False, "666666")


def style(doc: Document) -> None:
    sec = doc.sections[0]
    for side in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(sec, side, Inches(0.85))
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)


def method_summary_table() -> pd.DataFrame:
    return pd.DataFrame(
        [
            ["TotalForwardAnchor", 24.555, 45.963, 51.165, "总场锚点 HMM"],
            ["AxisAllMidGate", 17.586, 90.415, 107.305, "轴校准 HMM，长尾较大"],
            ["IMUProgressClosest_TotalVsAxis", 13.844, 27.533, 42.638, "当前最好，全段口径"],
            ["IMUProgressClosest，剔除严重真值异常", 15.681, 22.897, 29.439, "能力口径"],
        ],
        columns=["方法", "中位误差/m", "平均误差/m", "RMSE/m", "说明"],
    )


def build() -> None:
    doc = Document()
    style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("无轮速计铁路地磁定位：IMU 弱进度门控集成更新报告")
    set_font(r, 19, True, "0B2545")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("4.14 建图，5.13 跨日验证 | 2026-06-09")
    set_font(r, 10.5, False, "666666")

    heading(doc, "1. 当前结论", 1)
    paragraph(
        doc,
        "当前最好的无轮速计方案是：先分别生成总场锚点 HMM 候选轨迹和轴校准 HMM 候选轨迹，"
        "再用 INSPVAX 水平速度积分形成的弱进度一致性指标选择候选。该选择过程不使用 SPAN/GPGGA 真值。"
    )
    add_table(doc, method_summary_table(), font_size=8.3)

    heading(doc, "2. 方法说明", 1)
    bullets(
        doc,
        [
            "总场候选：使用 4.14 forward-only 锚点磁图、总场高通特征、单调 HMM/Viterbi，速度上界 vmax=1.2 m/s。",
            "轴候选：使用轴校准后的 X/Y/Total 高通特征和信息门控 HMM。它在部分反向段能避开总场重复特征假峰，但长尾更大。",
            "IMU 弱进度：对 INSPVAX 水平速度积分，得到整段位移量级；它不是轮速计，不逐步约束 HMM，只用于候选轨迹级选择。",
            "选择公式：compatibility = |log((candidate_progress + 10) / (imu_progress + 10))|，选择 compatibility 更小的候选。",
        ],
    )

    heading(doc, "3. 分段选择结果", 1)
    result_path = ROOT / "imu_progress_gated_ensemble" / "imu_progress_gated_ensemble_results.csv"
    results = pd.read_csv(result_path)
    selected = results[results["method"] == "IMUProgressClosest_TotalVsAxis"][
        [
            "segment_label",
            "direction",
            "chosen_candidate",
            "imu_distance_m",
            "total_progress_m",
            "axis_progress_m",
            "median_abs_error_m",
            "mean_abs_error_m",
            "rmse_m",
        ]
    ].copy()
    selected.columns = ["段", "方向", "选择候选", "IMU积分/m", "总场进度/m", "轴候选进度/m", "中位误差/m", "平均误差/m", "RMSE/m"]
    add_table(doc, selected, font_size=7.5)
    picture(
        doc,
        ROOT / "imu_progress_gated_ensemble" / "imu_progress_gated_ensemble_summary.png",
        "图 1 IMU 弱进度门控集成与两个单独候选方法的对比",
    )

    heading(doc, "4. 负结果与边界", 1)
    bullets(
        doc,
        [
            "延迟多假设 HMM 没有超过当前基线：最好 W=90 时中位误差 28.6 m，RMSE 73.7 m。",
            "仅改变候选累计似然的鲁棒统计方式也没有解决排序问题，说明需要独立证据，而不是同一总场分数的不同聚合。",
            "简单把 INSPVAX 速度作为逐步速度先验会恶化结果，因为速度积分在不同段上存在明显尺度不一致。",
            "1_seg03 同时存在严重 SPAN/GPGGA 距离轴跳变和重复磁特征假峰，后续应作为完整性检测和冷启动难例单独分析。",
        ],
    )

    heading(doc, "5. 论文创新点雏形", 1)
    bullets(
        doc,
        [
            "方向/质量感知锚点磁图：避免简单平均所有趟导致稳定磁特征被污染。",
            "无轮速单调 HMM：只依赖轨道一维约束、方向和速度上界。",
            "双候选磁匹配：总场候选提供姿态不敏感的稳健性，轴候选提供局部特征补充。",
            "IMU 弱进度门控：使用非轮速 IMU/INS 速度积分作为候选级完整性证据，而不是强里程计。",
        ],
    )

    heading(doc, "参考文献", 1)
    refs = [
        "Siebler et al., Train Localization with Particle Filter and Magnetic Field Measurements, FUSION 2018.",
        "Siebler et al., Magnetic Field Mapping of Railway Lines with Graph SLAM, FUSION 2024.",
        "Dieckow et al., Real-time rail vehicle localisation using spatially resolved magnetic field measurements, arXiv:2507.19327, 2025.",
        "WM-GFM: a novel geomagnetic feature matching positioning method with weak mileage aid, Measurement, 2026.",
    ]
    for ref in refs:
        paragraph(doc, ref, style="List Number")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
