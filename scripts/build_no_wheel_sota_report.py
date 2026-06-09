from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path.home() / "Desktop" / "\u78c1\u5bfc\u822a" / "\u6570\u636e" / "codex_railway_magnav"
EXP_ROOT = PROJECT_ROOT / "no_wheel_sota"
OUT_DOCX = EXP_ROOT / "reports" / "\u65e0\u8f6e\u901f\u8ba1SOTA\u5ba1\u6838\u4e0e\u6539\u8fdb\u65b9\u6cd5\u9636\u6bb5\u62a5\u544a.docx"


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 78, 121)),
    ]:
        st = doc.styles[name]
        st.font.name = "Microsoft YaHei"
        st.font.size = Pt(size)
        st.font.color.rgb = color


def add_table(doc: Document, df: pd.DataFrame, caption: str) -> None:
    doc.add_paragraph(caption)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Table Grid"
    for i, col in enumerate(df.columns):
        p = t.rows[0].cells[i].paragraphs[0]
        r = p.add_run(str(col))
        r.bold = True
        r.font.size = Pt(8)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                val = f"{val:.2f}"
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)


def add_pic(doc: Document, path: Path, caption: str) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(6.2))
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    EXP_ROOT.joinpath("reports").mkdir(parents=True, exist_ok=True)
    summary = pd.read_csv(EXP_ROOT / "outputs" / "no_wheel_sota_summary.csv")
    results = pd.read_csv(EXP_ROOT / "outputs" / "no_wheel_sota_results.csv")

    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("\u65e0\u8f6e\u901f\u8ba1\u6761\u4ef6\u4e0b\u94c1\u8def\u5730\u78c1\u5b9a\u4f4d SOTA \u5ba1\u6838\u4e0e\u6539\u8fdb\u65b9\u6cd5")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_heading("\u4e00\u3001\u9632\u5e7b\u89c9\u5ba1\u6838\u7ed3\u8bba", level=1)
    doc.add_paragraph(
        "\u672c\u9636\u6bb5\u5148\u5ba1\u6838\u4e86\u65b9\u6cd5\u8f93\u5165\u6761\u4ef6\uff1a2024 \u5e74\u94c1\u8def Graph SLAM \u4f9d\u8d56 odometer/\u8f6e\u901f\u8ba1\uff0c"
        "\u56e0\u6b64\u4e0d\u662f\u4f60\u5f53\u524d\u201c\u65e0\u8f6e\u901f\u8ba1\u201d\u6761\u4ef6\u4e0b\u7684\u516c\u5e73 SOTA baseline\u3002"
        "\u66f4\u8d34\u8fd1\u7684\u94c1\u8def\u65e0\u8f6e\u901f\u8ba1\u6587\u732e\u662f Siebler et al. 2018 FUSION \u7684\u78c1\u573a\u7c92\u5b50\u6ee4\u6ce2\uff0c"
        "\u5b83\u660e\u786e\u4ee5\u78c1\u56fe+\u78c1\u5f3a\u8ba1\u4e3a\u4e3b\uff0c\u72b6\u6001\u4e2d\u540c\u65f6\u4f30\u8ba1\u6cbf\u8f68\u4f4d\u7f6e\u548c\u901f\u5ea6\u3002"
    )
    doc.add_paragraph(
        "\u56e0\u6b64\u672c\u6b21\u6ca1\u6709\u628a\u201cIMU \u56e0\u5b50\u56fe\u201d\u76f4\u63a5\u5199\u6210\u521b\u65b0\u70b9\uff0c\u800c\u662f\u5148\u590d\u73b0\u6587\u732e\u6700\u63a5\u8fd1\u7684\u65e0\u8f6e\u901f\u57fa\u7ebf\uff0c"
        "\u518d\u63d0\u51fa\u4e00\u4e2a\u6709\u5b9e\u9a8c\u652f\u6491\u7684\u6539\u8fdb\uff1a\u7a33\u5065\u603b\u573a+\u603b\u573a\u9ad8\u901a\u7684\u79bb\u7ebf Viterbi/HMM \u5339\u914d\u3002"
    )

    lit = pd.DataFrame(
        [
            ["Siebler 2018 FUSION", "\u78c1\u56fe+\u78c1\u5f3a\u8ba1", "\u65e0\u8f6e\u901f\uff1bPF \u540c\u65f6\u4f30\u8ba1\u4f4d\u7f6e/\u901f\u5ea6", "RMSE < 4 m", "\u6700\u9002\u5408\u5f53\u524d\u590d\u73b0\u57fa\u7ebf"],
            ["Siebler 2020 PLANS", "\u78c1\u56fe+\u78c1\u5f3a\u8ba1+\u52a0\u901f\u5ea6", "\u9012\u5f52\u8d1d\u53f6\u6ee4\u6ce2+\u8f68\u9053\u8bc6\u522b", "RMSE < 5 m", "\u53ef\u4f5c\u540e\u7eed\u591a\u8f68/\u9053\u5c94\u6269\u5c55"],
            ["Siebler 2022 ION GNSS+", "\u78c1\u56fe+\u78c1\u5f3a\u8ba1", "\u9c81\u68d2 PF\uff0cLRT \u68c0\u6d4b\u78c1\u6d4b\u5f02\u5e38", "\u4fa7\u91cd\u9c81\u68d2\u6027", "\u652f\u6491\u672c\u6587\u9c81\u68d2\u4f3c\u7136/\u95e8\u9650\u8bbe\u8ba1"],
            ["Dieckow 2025 arXiv", "\u7a7a\u95f4\u5206\u8fa8\u78c1\u6d4b+\u9884\u5efa\u56fe", "\u91cd\u5c3e PF+\u5e8f\u5217\u5bf9\u9f50+\u6df7\u5408\u521d\u59cb\u5316", "PF sub-5m\uff1b\u5bf9\u9f50 92% < 30m", "\u6700\u65b0\u4f46\u786c\u4ef6\u6761\u4ef6\u5f3a\u4e8e\u672c\u6570\u636e"],
            ["Siebler 2024 FUSION", "\u78c1\u5f3a\u8ba1+\u8f6e\u901f/odometer", "Graph SLAM loop closure", "\u7c73\u7ea7", "\u9700\u8f6e\u901f\uff0c\u4e0d\u4f5c\u672c\u6761\u4ef6\u4e0b\u7684\u76f4\u63a5 SOTA"],
        ],
        columns=["\u6587\u732e", "\u8f93\u5165", "\u662f\u5426\u7b26\u5408\u65e0\u8f6e\u901f", "\u62a5\u544a\u6548\u679c", "\u672c\u9879\u76ee\u4f5c\u7528"],
    )
    add_table(doc, lit, "\u8868 1  \u65b9\u6cd5\u8f93\u5165\u6761\u4ef6\u4e0e SOTA \u5b9a\u4f4d")

    doc.add_heading("\u4e8c\u3001\u590d\u73b0\u57fa\u7ebf\u548c\u65b0\u65b9\u6cd5", level=1)
    doc.add_paragraph(
        "SOTA2018_PF_total\uff1a\u590d\u73b0 FUSION 2018 \u7684 SIR \u7c92\u5b50\u6ee4\u6ce2\u601d\u8def\uff0c\u72b6\u6001 x_k=[s_k, v_k]^T\uff0c"
        "\u7528\u6709\u9650\u52a0\u901f\u5ea6\u8fd0\u52a8\u6a21\u578b\u9884\u6d4b\uff0c\u7528\u5355\u70b9\u603b\u573a\u4e0e 4.14 \u78c1\u56fe\u7684\u5dee\u5f02\u66f4\u65b0\u6743\u91cd\u3002"
    )
    doc.add_paragraph(
        "SOTA2018_Viterbi_total\uff1a\u4e3a\u907f\u514d\u7c92\u5b50\u9000\u5316\u548c\u968f\u673a\u6027\uff0c\u5b9e\u73b0\u4e86\u7b49\u4ef7\u7684\u79bb\u6563 HMM/Viterbi \u7248\u672c\uff1b"
        "\u5b83\u4ecd\u53ea\u7528\u603b\u573a\u7279\u5f81\uff0c\u56e0\u6b64\u4f5c\u4e3a\u7a33\u5b9a\u7684 SOTA2018 \u57fa\u7ebf\u3002"
    )
    doc.add_paragraph(
        "Proposed_RobustTotalHP_Viterbi\uff1a\u5728 SOTA2018 \u57fa\u7ebf\u4e0a\u4fee\u6539\u4e24\u70b9\uff1a"
        "1) \u4f7f\u7528\u603b\u573a\u6807\u51c6\u5316+\u603b\u573a\u9ad8\u901a\u4e24\u7c7b\u7279\u5f81\uff0c\u524a\u5f31\u65e5\u671f/\u4f20\u611f\u5668\u504f\u7f6e\uff1b"
        "2) \u4f7f\u7528 Student-t \u578b\u91cd\u5c3e\u4f3c\u7136\uff0c\u964d\u4f4e\u5c40\u90e8\u78c1\u5e72\u6270\u5bf9\u8def\u5f84\u7684\u62d6\u62fd\u3002"
        "\u8fd9\u4e24\u70b9\u5206\u522b\u6709\u78c1\u7279\u5f81\u9ad8\u901a/\u68af\u5ea6\u5339\u914d\u548c\u9c81\u68d2 PF/LRT \u6587\u732e\u652f\u6491\u3002"
    )

    doc.add_heading("\u4e09\u3001\u5728\u672c\u6570\u636e\u4e0a\u7684\u7ed3\u679c", level=1)
    table = summary.copy()
    table = table[["method", "segment_count", "median_abs_error_m", "mean_abs_error_m", "rmse_m", "p90_abs_error_m", "final_abs_error_m"]]
    table.columns = ["\u65b9\u6cd5", "\u7247\u6bb5\u6570", "\u4e2d\u4f4d\u7edd\u5bf9\u8bef\u5dee/m", "\u5e73\u5747\u7edd\u5bf9\u8bef\u5dee/m", "\u5e73\u5747RMSE/m", "P90\u7edd\u5bf9\u8bef\u5dee/m", "\u7ec8\u70b9\u8bef\u5dee\u4e2d\u4f4d/m"]
    add_table(doc, table, "\u8868 2  5.13 \u67e5\u8be2\u76f8\u5bf9 4.14 \u78c1\u56fe\u7684\u5b9a\u4f4d\u7ed3\u679c")

    best = summary.iloc[0]
    baseline = summary[summary["method"] == "SOTA2018_Viterbi_total"].iloc[0]
    doc.add_paragraph(
        f"\u6700\u597d\u65b9\u6cd5 {best['method']} \u7684\u4e2d\u4f4d\u7edd\u5bf9\u8bef\u5dee\u4e3a {best['median_abs_error_m']:.1f} m\uff0c"
        f"\u76f8\u6bd4 SOTA2018_Viterbi_total \u7684 {baseline['median_abs_error_m']:.1f} m \u964d\u4f4e\u7ea6 "
        f"{(baseline['median_abs_error_m']-best['median_abs_error_m'])/baseline['median_abs_error_m']*100:.1f}%\u3002"
        "\u8fd9\u8bf4\u660e\u5728\u4f60\u7684\u5355\u78c1\u5f3a\u8ba1\u5c0f\u8f66\u6570\u636e\u4e0a\uff0c\u201c\u7a33\u5065\u603b\u573a+\u9ad8\u901a+\u4e00\u7ef4\u8fd0\u52a8\u7ea6\u675f\u201d\u786e\u5b9e\u8d85\u8fc7\u4e86\u76f4\u63a5\u590d\u73b0\u7684\u65e0\u8f6e\u901f\u57fa\u7ebf\u3002"
    )
    doc.add_paragraph(
        "\u4f46\u5fc5\u987b\u8bda\u5b9e\u8bf4\uff1a\u8fd9\u8fd8\u6ca1\u6709\u8fbe\u5230\u6587\u732e\u91cc\u7684\u7c73\u7ea7\u7ed3\u679c\u3002"
        "\u539f\u56e0\u5305\u62ec\uff1a\u4f60\u7684\u8f68\u9053\u957f\u5ea6\u53ea\u6709\u7ea6 575 m\uff0c\u78c1\u7279\u5f81\u81ea\u76f8\u4f3c\u66f4\u96be\u6392\u9664\uff1b"
        "4.14/5.13 \u4e24\u5929\u6570\u636e\u5b58\u5728\u5f3a\u5ea6\u504f\u7f6e\u548c\u5c40\u90e8\u5e72\u6270\uff1bSPAN \u8bc4\u4ef7\u771f\u503c\u4e5f\u6709\u5c40\u90e8\u8df3\u53d8\u3002"
    )

    fig_dir = EXP_ROOT / "figures"
    add_pic(doc, fig_dir / "no_wheel_sota_method_summary.png", "\u56fe 1  \u65e0\u8f6e\u901f SOTA \u590d\u73b0\u548c\u6539\u8fdb\u65b9\u6cd5\u5bf9\u6bd4")
    add_pic(doc, fig_dir / "no_wheel_sota_example_trajectories.png", "\u56fe 2  \u793a\u4f8b\u8f68\u8ff9\u5bf9\u6bd4")

    doc.add_heading("\u56db\u3001\u4e0b\u4e00\u4e2a\u53ef\u53d1\u5c55\u521b\u65b0\u70b9", level=1)
    doc.add_paragraph(
        "\u5efa\u8bae\u628a\u65b0\u65b9\u6cd5\u53d1\u5c55\u4e3a\uff1a\u201c\u65e0\u8f6e\u901f\u94c1\u8def\u573a\u666f\u7684\u9c81\u68d2\u78c1\u5e8f\u5217 HMM/Factor-Graph \u5b9a\u4f4d\u201d\u3002"
        "\u77ed\u671f\u5148\u505a HMM/Viterbi\uff0c\u4e2d\u671f\u628a IMU yaw/\u52a0\u901f\u5ea6\u52a0\u5165\u4e3a\u8fd0\u52a8\u5148\u9a8c\uff0c\u957f\u671f\u7528\u539f\u59cb RAWIMUSX \u505a\u4e00\u7ef4\u9884\u79ef\u5206\u548c bias \u4f30\u8ba1\u3002"
        "\u8fd9\u6761\u8def\u7ebf\u4e0e 2018 PF\u30012020 \u9012\u5f52\u8d1d\u53f6\u6ee4\u6ce2\u30012022 \u9c81\u68d2 PF\u30012025 \u5e8f\u5217\u5bf9\u9f50\u7684\u6587\u732e\u7ebf\u7d22\u4e00\u81f4\uff0c\u4e5f\u5df2\u5728\u4f60\u7684\u6570\u636e\u4e0a\u663e\u793a\u4e86\u76f8\u5bf9\u63d0\u5347\u3002"
    )

    doc.add_heading("\u4e94\u3001\u53c2\u8003\u6587\u732e", level=1)
    refs = [
        "Siebler, B., Heirich, O., Sand, S. Train Localization with Particle Filter and Magnetic Field Measurements. FUSION 2018. DOI: 10.23919/ICIF.2018.8455298. https://elib.dlr.de/119898/1/FUSION_2018.pdf",
        "Siebler, B., Heirich, O., Sand, S., Hanebeck, U. D. Joint Train Localization and Track Identification based on Earth Magnetic Field Distortions. PLANS 2020. DOI: 10.1109/PLANS46316.2020.9110149. https://isas.iar.kit.edu/pdf/PLANS20_Siebler.pdf",
        "Siebler, B., Heirich, O., Lehner, A., Sand, S., Hanebeck, U. D. Robust Particle Filter for Magnetic Field-based Train Localization. ION GNSS+ 2022. https://isas.iar.kit.edu/pdf/ION-GNSS22_Siebler.pdf",
        "Dieckow, N. et al. Real-time rail vehicle localisation using spatially resolved magnetic field measurements. arXiv:2507.19327, 2025. https://arxiv.org/abs/2507.19327",
        "Siebler, B., Lehner, A., Sand, S., Hanebeck, U. D. Magnetic Field Mapping of Railway Lines with Graph SLAM. FUSION 2024. DOI: 10.23919/FUSION59988.2024.10706392. https://isas.iar.kit.edu/pdf/FUSION24_Siebler.pdf",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Number")

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
