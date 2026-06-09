from __future__ import annotations

from pathlib import Path
import json
import math
import textwrap

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\m1352\Documents\railway_magnav")
OUT_DIR = ROOT / "reports" / "supervisor_brief_20260609"
OUT_DOCX = OUT_DIR / "铁路地磁定位阶段性技术路线与方向建议_20260609.docx"
OUT_MD = OUT_DIR / "铁路地磁定位阶段性技术路线与方向建议_20260609.md"


FIG_PROGRESS = ROOT / "progress_margin_selector_experiment" / "progress_margin_selector_summary.png"
FIG_TRAJ = ROOT / "progress_margin_selector_experiment" / "progress_margin_selector_selected_trajectories.png"
FIG_ENDPOINT = ROOT / "endpoint_prior_hmm_experiment" / "endpoint_prior_hmm_summary.png"
FIG_ROUTE = OUT_DIR / "technical_route_flow.png"
FIG_SOTA = OUT_DIR / "sota_metric_comparison.png"


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def wrap_text(text: str, width_chars: int) -> list[str]:
    out: list[str] = []
    for para in text.split("\n"):
        if not para:
            out.append("")
            continue
        out.extend(textwrap.wrap(para, width=width_chars, break_long_words=True, replace_whitespace=False))
    return out


def draw_wrapped(draw: ImageDraw.ImageDraw, xy, text: str, fnt, fill, width_chars: int, line_gap: int = 6):
    x, y = xy
    for line in wrap_text(text, width_chars):
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap
    return y


def create_route_figure(path: Path) -> None:
    w, h = 1800, 900
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_f = font(42, True)
    h_f = font(28, True)
    body_f = font(21)
    small_f = font(20)
    blue = (46, 116, 181)
    dark = (31, 58, 95)
    gray = (242, 244, 247)
    border = (185, 194, 205)
    d.text((60, 35), "当前铁路地磁定位技术路线", font=title_f, fill=dark)
    boxes = [
        ("数据对齐与坐标化", "SPAN/GPGGA 真值与磁强计时间对齐；沿轨方向投影为一维距离 s；按 0.5 m 网格建图。"),
        ("磁图构建", "总场高通特征为主；向量轴只做经验校准和辅助候选；多趟鲁棒融合参考磁图。"),
        ("候选匹配", "NCC/MSD/DTW 作基线；HMM/Viterbi 建模位置连续性；多 vmax 保留进度假设。"),
        ("弱 IMU 约束", "不把 IMU 当轮速计；只用整段进度一致性、速度软约束和候选筛选。"),
        ("完整性判断", "局部唯一性、top-k 间隔、Viterbi margin、进度残差、方法分歧共同决定可信度。"),
        ("评估协议", "全原始区段 + 单程单调区段同时报告；指标含中位、均值、RMSE、P90、终点误差。"),
    ]
    x0, y0 = 55, 120
    box_w, box_h = 530, 190
    gap_x, gap_y = 45, 55
    for i, (head, body) in enumerate(boxes):
        row, col = divmod(i, 3)
        x = x0 + col * (box_w + gap_x)
        y = y0 + row * (box_h + gap_y)
        d.rounded_rectangle((x, y, x + box_w, y + box_h), radius=18, fill=gray, outline=border, width=3)
        d.text((x + 25, y + 22), head, font=h_f, fill=blue)
        draw_wrapped(d, (x + 25, y + 70), body, body_f, (25, 25, 25), 21, 6)
    # arrows
    for row in range(2):
        for col in range(2):
            x = x0 + col * (box_w + gap_x) + box_w
            y = y0 + row * (box_h + gap_y) + box_h // 2
            d.line((x + 10, y, x + gap_x - 10, y), fill=blue, width=5)
            d.polygon([(x + gap_x - 10, y), (x + gap_x - 28, y - 12), (x + gap_x - 28, y + 12)], fill=blue)
    d.text((60, h - 70), "注意：反向尾段处理属于数据质检/区段定义，不作为方法创新。", font=small_f, fill=(120, 80, 0))
    img.save(path)


def create_sota_figure(path: Path) -> None:
    data = [
        ("Graph SLAM loop\nclosure 2024", 0.45, "相对位置 RMSE"),
        ("Snapshot/SLAC\n2025", 1.0, "RMSE <1 m"),
        ("Train PF\n2018", 3.84, "沿轨 RMSE"),
        ("Rail PF warm\nstart 2025", 5.0, "sub-5 m"),
        ("本项目最优\n清洗后", 39.0, "RMSE"),
        ("本项目最优\n原始全段", 40.1, "RMSE"),
    ]
    w, h = 1550, 820
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_f = font(40, True)
    lab_f = font(22)
    small_f = font(19)
    axis_f = font(20)
    d.text((60, 35), "铁路地磁定位代表性指标对比（纵轴为对数尺度）", font=title_f, fill=(31, 58, 95))
    plot_left, plot_top, plot_w, plot_h = 130, 130, 1300, 500
    d.line((plot_left, plot_top + plot_h, plot_left + plot_w, plot_top + plot_h), fill=(80, 80, 80), width=2)
    d.line((plot_left, plot_top, plot_left, plot_top + plot_h), fill=(80, 80, 80), width=2)
    min_log, max_log = math.log10(0.3), math.log10(80)
    ticks = [0.5, 1, 2, 5, 10, 20, 50]
    for t in ticks:
        y = plot_top + plot_h - (math.log10(t) - min_log) / (max_log - min_log) * plot_h
        d.line((plot_left - 8, y, plot_left + plot_w, y), fill=(225, 225, 225), width=1)
        d.text((35, y - 12), f"{t:g} m", font=axis_f, fill=(80, 80, 80))
    bar_gap = 38
    bar_w = (plot_w - bar_gap * (len(data) + 1)) / len(data)
    colors = [(90, 155, 212), (90, 155, 212), (90, 155, 212), (90, 155, 212), (200, 82, 82), (200, 82, 82)]
    for i, (name, value, note) in enumerate(data):
        x = plot_left + bar_gap + i * (bar_w + bar_gap)
        y = plot_top + plot_h - (math.log10(value) - min_log) / (max_log - min_log) * plot_h
        d.rounded_rectangle((x, y, x + bar_w, plot_top + plot_h), radius=8, fill=colors[i])
        d.text((x + 5, y - 32), f"{value:g} m", font=lab_f, fill=(40, 40, 40))
        tx = x - 5
        ty = plot_top + plot_h + 18
        for line in name.split("\n"):
            d.text((tx, ty), line, font=small_f, fill=(40, 40, 40))
            ty += 24
        d.text((tx, ty + 4), note, font=small_f, fill=(95, 95, 95))
    d.text((60, 735), "解读：本项目当前证明了磁特征可用，但与列车平台/里程计辅助/长线路数据的 SOTA 相比仍差一个数量级。", font=lab_f, fill=(120, 80, 0))
    img.save(path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_east_asia_font(run, name: str = "Microsoft YaHei") -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    set_east_asia_font(r)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=8)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    return table


def add_paragraph(doc: Document, text: str, style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def style_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    styles = doc.styles
    for name in ["Normal", "Body Text", "List Bullet", "List Number"]:
        if name in styles:
            st = styles[name]
            st.font.name = "Calibri"
            st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            st.font.size = Pt(11)
            st.paragraph_format.space_after = Pt(6)
            st.paragraph_format.line_spacing = 1.10
    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)


def add_picture_if_exists(doc: Document, path: Path, caption: str, width_in: float = 6.2):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(90, 90, 90)


def build_markdown() -> str:
    return """# 铁路地磁定位阶段性技术路线与方向建议

日期：2026-06-09  
项目：良陈铁路约 700 m 小车地磁采集，4.14 建图，5.13 跨日验证

## 核心结论

1. 当前数据证明铁路沿线地磁特征具有可重复性，但目前还不能支撑“达到铁路地磁定位 SOTA”的结论。
2. 当前最优方法为 Progress-Margin Selector：多 vmax 总场 HMM 候选 + 轴校准候选 + 弱 IMU 进度与 Viterbi margin 筛选。5 个可用 5.13 段的原始全段 RMSE 为 40.1 m，中位误差 13.8 m；单程单调区段 RMSE 为 39.0 m，平均终点误差 16.7 m。
3. 与代表性铁路 SOTA（1 m 到 5 m 级）差距明显，主要差在数据平台、里程/速度约束、地图规模、多日数据量、在线完整性判断，而不是“地磁不可用”。
4. 小车采集铁路地磁数据可行，但当前 700 m、两天、无轮速计的数据只能支撑可行性验证和方法雏形；若要发三区论文，建议至少补到 4 天、20 条以上单程、3 个以上跨日验证组合。
5. 后续方向建议：短期继续铁路，作为受约束的一维场景把方法做实；同时准备城市车辆平台作为备选或下一阶段扩展。若导师更看重应用面和数据规模，城市车辆平台更有发展空间；若更看重可控性和快速形成论文闭环，铁路方向仍值得再做一轮高质量采集。

## 最主要参考文献

最主要参考文献建议放 Siebler, Heirich, Sand, *Train Localization with Particle Filter and Magnetic Field Measurements*, FUSION 2018。原因是它和本项目最接近：都是铁路沿轨一维地磁地图定位，都用已有磁图和运动模型约束位置估计，并且其 SIR 粒子滤波报告沿轨 RMSE 3.84 m，是当前“无轮速/弱运动约束铁路地磁定位”最重要的公平基线。

强 SOTA 对标文献再补两篇：FUSION 2024 *Magnetic Field Mapping of Railway Lines with Graph SLAM*，以及 2025 arXiv *Real-time rail vehicle localisation using spatially resolved magnetic field measurements*。

## 当前指标最好方法怎么做

当前指标最好的方法叫 Progress-Margin Selector。它不是单独用一条磁曲线硬匹配，而是先生成多条候选轨迹，再根据弱 IMU 进度和匹配置信度选择最可信的一条。

1. 用 4.14 数据建立参考磁图，主特征用总场高通值，原因是总场不受小车朝向翻转影响，比三轴更稳定。
2. 对 5.13 每个查询段做 HMM/Viterbi 匹配。HMM 状态是沿轨位置 s_t，转移约束是相邻时刻不能跳太远，观测似然来自当前位置磁场与磁图是否相似。
3. 不只跑一个速度上限，而是分别跑 vmax=1.0、1.2、1.4 m/s，得到多条总场候选轨迹，用来处理小车速度不稳定的问题。
4. 另外跑一条轴校准候选。三轴磁数据有时能补充总场，但跨日不够稳定，所以只作为候选，不作为默认主方法。
5. 对每条候选计算 candidate_progress，再与 INSPVAX 积分得到的 imu_progress 比较：progress_compat = |log((candidate_progress + 10)/(imu_progress + 10))|。数值越小，说明这条候选轨迹的总进度越像真实运动。
6. 在总场候选里，先保留 progress_compat 接近最优的候选，再看 Viterbi final_score_margin；如果仍然相近，偏向更大的 vmax，避免末端跟不上。
7. 只有当轴校准候选的 Viterbi margin 足够高，并且进度一致性明显优于总场候选时，才切换到轴候选。
8. SPAN/GPGGA 真值只用于最后算误差，不参与候选选择。
"""


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_route_figure(FIG_ROUTE)
    create_sota_figure(FIG_SOTA)
    OUT_MD.write_text(build_markdown(), encoding="utf-8")

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("铁路地磁定位阶段性技术路线与方向建议")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 58, 95)
    set_east_asia_font(r)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("用于导师汇报 | 2026-06-09 | 良陈铁路约 700 m 小车数据")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(90, 90, 90)
    set_east_asia_font(r)

    doc.add_heading("1. 一句话结论", level=1)
    add_paragraph(
        doc,
        "当前结果说明铁路地磁特征可以用于跨日匹配，但现阶段还只是“可行性 + 离线/延迟匹配雏形”，距离铁路地磁定位 SOTA 的米级精度和可发表三区期刊的严谨验证还有明显差距；建议先补采高质量多日数据并发展多假设 PF/HMM + 弱 IMU 进度约束 + 完整性评分方法，再决定是否继续深挖铁路或转向城市车辆平台。",
    )

    doc.add_heading("1.1 最主要参考文献", level=2)
    add_paragraph(
        doc,
        "本阶段最主要的参考文献建议明确写为：Siebler, Heirich, Sand, Train Localization with Particle Filter and Magnetic Field Measurements, FUSION 2018。",
    )
    add_paragraph(
        doc,
        "选择它作为主参考的原因是：它与本项目问题最接近，都是铁路沿轨一维地磁地图定位，都依赖已有磁图、磁强计观测和运动模型约束来估计沿轨位置；其 SIR 粒子滤波报告沿轨 RMSE 3.84 m、最大误差 43.48 m，是目前与“无轮速/弱运动约束铁路地磁定位”最接近的公平基线。",
    )
    add_paragraph(
        doc,
        "另外两篇作为强 SOTA 对标：FUSION 2024 的 Magnetic Field Mapping of Railway Lines with Graph SLAM 代表“里程计 + 多趟磁 loop closure + 图优化”的强建图路线；2025 arXiv 的 Real-time rail vehicle localisation using spatially resolved magnetic field measurements 代表“top-k 空间序列初始化 + 重尾粒子滤波”的最新实时定位路线。",
    )

    doc.add_heading("2. 数据与问题定义", level=1)
    add_bullets(
        doc,
        [
            "数据来源：4.14 与 5.13 两次铁路小车采集，线路约 700 m，包含 SPAN/GPGGA 位置真值、磁强计三轴与总场数据。",
            "当前处理：将 GNSS 投影到沿轨一维距离 s，按 0.5 m 网格构建磁图；4.14 作为参考图，5.13 作为跨日查询验证。",
            "约束条件：无轮速计；IMU/INSPVAX 可用但存在漂移与尺度不稳定，因此仅作为弱进度约束，不等价于里程计。",
            "评估口径：同时报告原始全段指标和单程单调区段指标；反向尾段只作为数据质检/区段定义，不写成方法创新。",
        ],
    )

    doc.add_heading("3. 目前技术路线", level=1)
    add_picture_if_exists(doc, FIG_ROUTE, "图 1  当前铁路地磁定位技术路线")
    add_paragraph(doc, "核心流程如下。")
    add_numbered(
        doc,
        [
            "时间与坐标对齐：磁强计系统时间与 SPAN/GPGGA 时间对齐；将经纬度/平面坐标投影到轨道主方向，得到一维距离 s。",
            "磁图构建：以总场高通特征为主，三轴向量经轨道坐标/经验轴校准后作为辅助；多趟数据按距离网格鲁棒融合。",
            "基线匹配：滑窗 NCC、MSD、DTW 用于验证磁特征可重复性和粗匹配能力。",
            "HMM/Viterbi 匹配：状态为沿轨位置 s_t，转移由方向和最大速度约束，观测似然来自磁特征残差。",
            "弱 IMU 进度筛选：比较候选轨迹进度与 INSPVAX 积分进度的一致性，用作候选筛选，不把 IMU 当轮速计。",
            "完整性评分：后续应加入 top-k 间隔、局部唯一性、Viterbi margin、方法分歧、IMU 进度残差，判断是否可信。",
        ],
    )
    add_paragraph(doc, "关键公式：")
    add_bullets(
        doc,
        [
            "NCC：rho = sum((q-mean(q))(m-mean(m))) / (||q-mean(q)|| ||m-mean(m)||)，用于滑窗相似度匹配。",
            "MSD：MSD = mean((q - m)^2)，用于粗-精匹配中的距离残差评价。",
            "HMM 观测似然：log p(z_t|s_t) 约等于 -0.5 * r_t(s_t)^2 / sigma^2，r 为磁特征残差；实际使用鲁棒似然降低异常点影响。",
            "弱进度一致性：progress_compat = |log((candidate_progress + 10) / (imu_progress + 10))|，用于比较候选轨迹和 IMU 粗进度。",
            "误差：e_t = s_hat_t - s_true_t；报告 median |e|、mean |e|、RMSE、P90 |e|、endpoint |e_T|。",
        ],
    )

    doc.add_heading("4. 当前实验指标", level=1)
    doc.add_heading("4.1 当前指标最好的方法：Progress-Margin Selector", level=2)
    add_paragraph(
        doc,
        "简单说，Progress-Margin Selector 不是直接拿一条磁曲线去硬匹配，而是先生成多条可能的沿轨轨迹，再用弱 IMU 进度和匹配置信度挑出最可信的一条。它的目标是解决短铁路区段里“很多位置的磁特征看起来相似”的问题。",
    )
    add_numbered(
        doc,
        [
            "建立参考磁图：用 4.14 数据生成沿轨 0.5 m 网格磁图。主特征采用总场高通值，因为总场不受小车朝向翻转影响，跨日稳定性比三轴更好；三轴只作为辅助候选。",
            "生成总场 HMM 候选：对 5.13 每个查询段运行 HMM/Viterbi。状态为沿轨位置 s_t，观测似然表示当前位置磁场与参考磁图的相似程度，转移约束表示相邻时刻不能跳太远。",
            "保留多个速度假设：分别用 vmax = 1.0、1.2、1.4 m/s 生成候选轨迹。这样做是因为小车速度不稳定，固定一个速度上限容易导致末端跟不上或局部跑偏。",
            "生成轴校准候选：额外运行一条基于经验轴校准的候选。三轴有时能补充总场，但跨日安装/朝向不够稳定，所以它只参与竞争，不作为默认主方法。",
            "计算弱 IMU 进度一致性：对每条候选轨迹计算 candidate_progress = |s_hat_end - s_hat_start|，再与 INSPVAX 速度积分得到的 imu_progress 比较。使用 progress_compat = |log((candidate_progress + 10)/(imu_progress + 10))|，数值越小表示候选轨迹的总进度越像真实运动。",
            "总场候选筛选规则：先保留 progress_compat 与最优值相差不超过 0.04 的候选，再优先选择 final_score_margin 接近最优的候选；如果仍然相近，偏向更大的 vmax，避免末端滞后。",
            "是否切换到轴候选：只有当轴候选 final_score_margin >= 5，并且轴候选的 progress_compat 至少比已选总场候选好 0.04 时，才切换到轴候选。",
            "误差评估：SPAN/GPGGA 真值只用于最后计算误差，不参与候选选择；因此这个方法不是用真值挑结果。",
        ],
    )
    add_paragraph(
        doc,
        "这个方法有效的直观原因是：总场特征跨日最稳，多 vmax 候选能覆盖速度变化，弱 IMU 进度能排除一部分“磁场看起来像但进度不合理”的错误候选，Viterbi margin 则提供匹配置信度。它仍然是离线/延迟选择器，不是最终实时导航算法；后续要把这个思想升级为多假设 PF/HMM。",
    )

    doc.add_heading("4.2 指标汇总", level=2)
    add_table(
        doc,
        ["方法/评价集", "段数", "中位误差", "均值误差", "RMSE", "P90", "平均终点误差", "最大终点误差"],
        [
            ["FixedTotal HMM", "5", "24.6 m", "46.0 m", "51.2 m", "71.8 m", "50.6 m", "93.6 m"],
            ["ProgressMarginSelector 原始全段", "5", "13.8 m", "25.2 m", "40.1 m", "56.6 m", "24.6 m", "71.1 m"],
            ["ProgressMarginSelector 单程单调", "5", "13.8 m", "24.3 m", "39.0 m", "-", "16.7 m", "31.9 m"],
            ["剔除 1_seg03 + 单程单调", "4", "14.1 m", "18.9 m", "24.9 m", "-", "18.3 m", "31.9 m"],
            ["端点先验 HMM + 轻 IMU 速度", "5", "29.0 m", "53.5 m", "61.1 m", "89.0 m", "31.6 m", "71.1 m"],
        ],
        widths=[1.7, 0.45, 0.75, 0.75, 0.7, 0.7, 0.9, 0.9],
    )
    add_picture_if_exists(doc, FIG_PROGRESS, "图 2  当前候选方法误差汇总", 6.2)
    add_picture_if_exists(doc, FIG_TRAJ, "图 3  ProgressMarginSelector 分段轨迹结果", 6.2)
    add_picture_if_exists(doc, FIG_ENDPOINT, "图 4  端点先验 HMM 实验：端点先验本身不足以解决重复特征假匹配", 6.2)
    add_paragraph(
        doc,
        "组内可重复性补充：4.14 留一法显示，直接同距相关性部分段较差，但允许距离微小弯曲的 DTW 相关性普遍可达到约 0.82-0.94，说明磁特征本身可重复，主要问题在进度/速度模型、短线路重复特征和跨日传感器坐标不一致。",
    )

    doc.add_heading("5. 与代表性 SOTA 的差距", level=1)
    add_picture_if_exists(doc, FIG_SOTA, "图 5  代表性铁路地磁定位指标对比")
    add_table(
        doc,
        ["年份", "代表性工作", "平台/假设", "核心方法", "报告指标", "与本项目关系"],
        [
            ["2018", "Train Localization with Particle Filter and Magnetic Field Measurements", "真实列车；已有磁图；列车运动模型", "SIR 粒子滤波", "沿轨 RMSE 3.84 m，最大误差 43.48 m", "最接近无轮速公平基线，但其数据平台和线路尺度更强"],
            ["2022", "Robust Particle Filter for Magnetic field-based Train Localization", "铁路磁图；考虑故障/异常测量", "鲁棒 PF + 故障检测", "强调异常测量下鲁棒性提升", "支持我们做鲁棒似然和完整性门限"],
            ["2024", "Magnetic Field Mapping of Railway Lines with Graph SLAM", "磁强计 + 里程计；多趟运行", "局部磁签名相关形成 loop closure，pose graph 优化", "磁 loop closure 相对位置 RMSE 0.45 m，里程累计误差约束到米级", "强 SOTA 参考，但依赖里程计，不是无轮速公平对比"],
            ["2025", "Real-time rail vehicle localisation using spatially resolved magnetic field measurements", "长线路真实轨道；空间磁图", "重尾 PF + stateless sequence alignment", "21.6 km 上 warm-start PF sub-5 m；冷启动 top-1 30 m 内 92%", "最新强参考；说明 top-k 初始化 + PF 是重要路线"],
            ["2025", "Snapshot Estimator with Uncalibrated Magnetometers", "未标定磁强计；虚拟磁阵列/里程辅助", "SLAC/最大似然联合估计位置与标定", "8 km 测试线 2 Hz 下 RMSE <1 m", "启发我们做局部轴/标定自适应，但单独使用会过拟合错误候选"],
        ],
        widths=[0.45, 1.35, 1.25, 1.2, 1.25, 1.5],
    )

    doc.add_heading("6. 为什么目前达不到 SOTA", level=1)
    add_bullets(
        doc,
        [
            "数据平台差距：SOTA 多为真实列车长距离运行，运动模型更稳定；我们是人工推小车，速度慢、停顿/掉头/反向尾段更复杂。",
            "里程约束差距：Graph SLAM 和 SLAC 类方法通常有里程计、虚拟阵列或更强运动先验；我们没有轮速计，IMU 只能弱约束，不能长期积分成可靠里程。",
            "线路长度与唯一性差距：700 m 短线路中存在重复磁签名，滑窗匹配容易出现多个相似候选；长线路反而能通过更长序列和更多上下文消除歧义。",
            "地图数据量差距：当前本质上是一天建图、一天验证；缺少多日、多温度、多速度、多载体姿态下的稳定性验证。",
            "坐标/安装一致性差距：三轴磁数据受小车朝向、传感器安装、车体铁磁环境影响；跨日向量轴不完全一致，所以目前总场比三轴更稳。",
            "算法成熟度差距：当前最优结果是离线/延迟候选选择，固定滞后和弱里程在线实验明显变差，说明还没形成可宣称实时导航的闭环算法。",
            "真值与分段问题：GPGGA 总体可靠，但个别段投影距离存在跳变或非单程运动；这需要透明报告，而不能只展示清洗后的好结果。",
        ],
    )

    doc.add_heading("7. 小车采集铁路地磁是否不可行", level=1)
    add_paragraph(
        doc,
        "不是不可行。现有数据已经证明铁路磁特征具有跨趟、跨日可重复性：4.14 组内 DTW 相关性普遍较高，5.13 对 4.14 的最优离线匹配中位误差达到 13.8 m。但当前小车实验更适合证明“铁路地磁特征存在、可匹配”，还不足以证明“无轮速实时铁路导航可以达到米级”。",
    )
    add_paragraph(doc, "小车平台可行的前提是把采集协议做扎实：")
    add_bullets(
        doc,
        [
            "固定传感器安装高度、朝向和远离车体铁磁干扰的位置；",
            "每趟记录起点、终点、是否掉头、停车、异常操作；",
            "增加独立里程/速度参考，哪怕是低成本轮编码器、视觉里程、RTK/GNSS 后处理或人工里程标，也能显著改善研究可信度；",
            "采集更多天和更多往返，覆盖不同速度、不同天气/温度和不同操作人员；",
            "尽量延长线路或选择磁特征更丰富、长度更长的区段，提高全局唯一性。",
        ],
    )

    doc.add_heading("8. 后续应该继续铁路还是转城市车辆", level=1)
    add_table(
        doc,
        ["方向", "优势", "风险", "建议"],
        [
            ["继续铁路", "一维约束强、问题边界清楚、已有数据和处理链、容易做完整误差评估", "SOTA 已较强；短线路重复特征明显；无轮速很吃亏；采集资源受限", "短期继续一轮高质量采集，把方法和论文闭环做出来"],
            ["城市车辆地磁", "应用面更广；导师建议方向；可与 GNSS/IMU/道路地图融合；路线更长，数据更丰富", "变成二维/车道级问题；环境变化和车辆铁磁干扰更复杂；建图成本更大", "中长期可转，尤其适合做 GNSS 受限城市峡谷融合定位"],
        ],
        widths=[1.0, 1.8, 1.8, 1.8],
    )
    add_paragraph(
        doc,
        "我的建议是“两步走”：不要立刻放弃铁路，因为当前已有数据和可重复性证据，适合快速形成一篇受控场景下的方法论文；但如果下一轮采集仍无法获得更长线路、多日数据或可靠速度/里程约束，则应尽早转向车辆平台。车辆方向更容易扩展为 GNSS/IMU/磁图/道路约束融合，论文空间更宽。",
    )

    doc.add_heading("9. 下一阶段工作计划", level=1)
    add_numbered(
        doc,
        [
            "补采数据：至少 4 天、20 条以上单程、3 个以上跨日验证组合；保留完整元数据和异常记录。",
            "建立强基线：总场 NCC/MSD/DTW、HMM/Viterbi、PF、MSD+ICCP 粗精匹配全部统一评价。",
            "发展主方法：top-k 磁候选生成 + 多假设 PF/HMM，状态包含位置 s、速度/进度尺度、方向模式和特征模型候选。",
            "加入弱 IMU 与完整性：IMU 只作为软进度约束；建立可信度评分，能拒绝局部磁特征不唯一的区段。",
            "报告规范：同时报告全原始区段、单程单调区段、剔除明显真值异常后的敏感性分析，不能只保留好看的段。",
            "决策节点：如果新增数据后跨日 RMSE 仍高于 30-40 m，且在线 PF 无法压到 20 m 内，应转向城市车辆平台或加入轮速/视觉里程约束。",
        ],
    )

    doc.add_heading("10. 参考文献与资料来源", level=1)
    refs = [
        "Siebler, B., Heirich, O., Sand, S. Train Localization with Particle Filter and Magnetic Field Measurements. FUSION 2018. https://elib.dlr.de/119898/1/FUSION_2018.pdf",
        "Siebler et al. Robust Particle Filter for Magnetic field-based Train Localization. ION GNSS+ 2022. https://www.ion.org/publications/abstract.cfm?articleID=18536",
        "Siebler et al. Magnetic Field Mapping of Railway Lines with Graph SLAM. FUSION 2024. https://isas.iar.kit.edu/pdf/FUSION24_Siebler.pdf",
        "Dieckow et al. Real-time rail vehicle localisation using spatially resolved magnetic field measurements. arXiv:2507.19327. https://arxiv.org/abs/2507.19327",
        "Siebler et al. Snapshot Estimator for Magnetic Field-based Train Localization with Uncalibrated Magnetometers. EUSIPCO 2025. https://eusipco2025.org/wp-content/uploads/pdfs/0002142.pdf",
        "WM-GFM: A novel geomagnetic feature matching positioning method with weak mileage aid. Measurement 257:118850. https://www.sciencedirect.com/science/article/abs/pii/S026322412600268X",
        "本项目结果文件：progress_margin_selector_experiment、turnaround_and_trim_diagnostic、endpoint_prior_hmm_experiment、latest_literature_aligned_experiments。",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_docx()
    print(OUT_DOCX)
    print(OUT_MD)
