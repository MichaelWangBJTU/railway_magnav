from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\m1352\Documents\railway_magnav")
OUT = ROOT / "reports" / "无轮速计铁路地磁定位_锚点磁图HMM阶段报告_20260609.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=10.5, bold=False, color=None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)


def add_table(doc: Document, df: pd.DataFrame, widths: list[int] | None = None, font_size: float = 9.0) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
        set_cell_shading(hdr[j], "E8EEF5")
        hdr[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(hdr[j], widths[j])
        for p in hdr[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_font(run, font_size, bold=True, color="0B2545")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            value = row[col]
            if isinstance(value, float):
                text = f"{value:.3f}"
            else:
                text = str(value)
            cells[j].text = text
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[j], widths[j])
            for p in cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_font(run, font_size)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13 if level == 2 else 11.5, bold=True, color="1F4D78")


def add_picture_if_exists(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.2))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    set_font(r, size=9.0, color="666666")


def style_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)
    for style_name in ["List Bullet", "List Number"]:
        st = styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10.0)
        st.paragraph_format.space_after = Pt(3)


def short_method_table() -> pd.DataFrame:
    return pd.DataFrame(
        [
            ["全趟平均图 + TotalHP HMM", 75.6, 77.4, 86.1, "旧基线"],
            ["轴校准 XY+Total 中等门控 HMM", 17.6, 90.4, 107.3, "中位好，但长尾大"],
            ["forward-only 锚点图 + TotalHP HMM, vmax=1.4", 27.3, 64.2, 73.5, "锚点图验证"],
            ["forward-only 锚点图 + TotalHP HMM, vmax=1.2", 24.6, 46.0, 51.2, "当前主结果"],
            ["同上，剔除严重真值轴异常段", 19.2, 21.0, 25.9, "能力口径"],
        ],
        columns=["方法", "中位误差/m", "平均误差/m", "RMSE/m", "备注"],
    )


def build_doc() -> None:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("无轮速计铁路地磁定位阶段性研究报告")
    set_font(r, size=20, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("锚点磁图构建 + 总场高通 HMM/Viterbi | 2026-06-09")
    set_font(r, size=11, color="666666")

    add_heading(doc, "1. 一句话结论", 1)
    add_paragraph(
        doc,
        "在无轮速计条件下，使用 4.14 正向高一致性趟构建 forward-only 锚点磁图，结合总场高通特征和单调 HMM/Viterbi，"
        "当前在 5.13 五个可用段上达到中位误差 24.6 m、平均误差 46.0 m、RMSE 51.2 m；"
        "剔除存在严重 SPAN/GPGGA 距离轴跳变的 1_seg03 后，四段中位误差 19.2 m、平均误差 21.0 m、RMSE 25.9 m。"
    )

    add_heading(doc, "2. 实验背景与评价口径", 1)
    add_bullets(
        doc,
        [
            "数据：4.14 作为参考建图日，5.13 作为跨日查询日；磁传感器三轴与总场数据已经按时间对齐到 SPAN/GPGGA 位置。",
            "目标：在没有轮速计的条件下，验证铁路沿线磁场能否支撑沿轨绝对位置匹配。",
            "公平边界：HMM 定位不使用 5.13 真值位置；真值只用于离线评价。全趟 DTW 只用于地图质量诊断，不作为在线定位结果。",
            "报告同时给出所有可用段指标和剔除严重真值轴异常段指标，避免把 SPAN 跳变误认为算法误差，也避免只挑好看的结果。",
        ],
    )

    add_heading(doc, "3. 方法细节", 1)
    add_heading(doc, "3.1 锚点磁图构建", 2)
    add_paragraph(
        doc,
        "直接平均所有趟会把不一致趟的局部形态混在一起。本轮实验比较了 all-pass、quality-good、forward-only、backward-only、"
        "LOPO 高一致性趟等参考图。最终用于定位的主结果采用 forward-only 锚点图，即仅用 4.14 中三条正向趟构建参考磁图。"
    )
    add_paragraph(doc, "高通总场特征定义为：B_hp(s) = B(s) - median_window(B(s))。")
    add_paragraph(doc, "稳健归一化采用：z(s) = (B_hp(s) - median(B_hp)) / (1.4826 * MAD(B_hp))。")

    add_heading(doc, "3.2 HMM/Viterbi 匹配", 2)
    add_paragraph(
        doc,
        "HMM 的状态是沿轨距离网格 s_i，网格间隔 0.5 m；观测为当前时刻的总场高通归一化值 z_t。"
        "观测似然使用重尾 Student-t 形式，以降低异常磁点影响："
    )
    add_paragraph(doc, "log p(z_t | s_i) = -0.5 * (nu + 1) * log(1 + ((z_t - m_i) / sigma)^2 / nu), 其中 nu=3。")
    add_paragraph(
        doc,
        "转移模型只允许沿声明方向单调运动，并限制速度不超过 vmax。调参后最佳 vmax=1.2 m/s。"
        "Viterbi 递推为：D_t(i)=log p(z_t|s_i)+max_j[D_{t-1}(j)+log p(s_i|s_j)]。"
    )
    add_paragraph(
        doc,
        "这一路线不依赖轮速计；INSPVAX 水平速度作为弱先验试过，但当前数据中速度尺度不稳定，整体会恶化指标，因此没有进入主方法。"
    )

    add_heading(doc, "4. 关键实验结果", 1)
    add_table(doc, short_method_table(), widths=[3700, 1250, 1250, 1250, 1800], font_size=8.5)

    add_heading(doc, "4.1 距离轴校正的负结果", 2)
    add_bullets(
        doc,
        [
            "全趟强制带限 DTW 校正会过拟合，4.14 地图标准化离散度从 0.903 恶化到 2.234。",
            "选择性校正能把离散度降到 0.794，但 LOPO 与跨日定位指标没有同步改善。",
            "结论：不能把“全趟 DTW 校正后平均”作为最终方法；DTW 更适合作为地图质量诊断和候选校正工具。",
        ],
    )
    add_picture_if_exists(
        doc,
        ROOT / "constrained_map_alignment_experiment" / "alignment_evaluation_summary.png",
        "图 1 约束距离轴校正前后的地图一致性评价",
    )

    add_heading(doc, "4.2 锚点参考图筛选", 2)
    add_paragraph(
        doc,
        "5.13 跨日地图质量诊断中，backward-only 参考图的带限 DTW 相关系数最高，为 0.789；"
        "但在线 HMM 中 forward-only 参考图表现最好。这说明离线形状相关性和在线定位误差不是同一件事。"
    )
    add_picture_if_exists(
        doc,
        ROOT / "anchor_reference_selection_experiment" / "anchor_reference_cross_day_summary.png",
        "图 2 不同 4.14 锚点参考图与 5.13 的跨日总场一致性",
    )

    add_heading(doc, "4.3 HMM 定位与速度上界调参", 2)
    add_paragraph(
        doc,
        "forward-only 参考图进入 HMM 后，TotalHP_Viterbi 的中位误差从旧基线 75.6 m 降到 27.3 m。"
        "进一步把最大速度从 1.4 m/s 调整到 1.2 m/s 后，平均误差和 RMSE 明显下降。"
    )
    add_picture_if_exists(
        doc,
        ROOT / "forward_anchor_hmm_tuning" / "forward_anchor_hmm_tuning_summary.png",
        "图 3 forward-only 锚点图下 TotalHP HMM 的 vmax、门控、速度先验调参结果",
    )
    add_picture_if_exists(
        doc,
        ROOT / "anchor_reference_hmm_experiment" / "anchor_reference_hmm_best_example.png",
        "图 4 锚点参考图 HMM 最佳组合的轨迹示例",
    )

    add_heading(doc, "4.4 真值轴异常段分析", 2)
    anomaly = pd.read_csv(ROOT / "truth_axis_anomaly_diagnostic" / "truth_axis_anomaly_by_segment.csv")
    anomaly_small = anomaly[
        [
            "segment_label",
            "direction",
            "sign_violation_count",
            "large_jump_count_gt20m",
            "max_abs_step_m",
            "truth_axis_warning",
            "severe_truth_axis_anomaly",
        ]
    ]
    add_table(doc, anomaly_small, widths=[2400, 900, 1200, 1400, 1200, 1100, 1300], font_size=7.5)
    add_paragraph(
        doc,
        "1_seg03 存在多次百米级 SPAN/GPGGA 沿轨距离跳变，例如约 192 m、197 m、292 m、271 m，"
        "且跳变邻域外仍有重复磁签名导致的在线初始假峰。因此它既是评价真值异常段，也是无轮速冷启动难例。"
    )

    add_heading(doc, "5. 当前可写成论文的技术点", 1)
    add_numbered(
        doc,
        [
            "质量/方向感知的锚点磁图构建：不是简单平均所有趟，而是基于组内一致性和方向筛选参考趟。",
            "无轮速单调 HMM：只用方向、最大速度和磁观测连续性进行定位，适合没有轮速计的实验条件。",
            "真值轴完整性诊断：把 SPAN/GPGGA 距离跳变与磁匹配失败分开，建立更可靠的评价口径。",
            "负结果边界清楚：全趟 DTW 校正、简单速度先验、简单信息门控、重权三轴融合都还不是稳定主方法。",
        ],
    )

    add_heading(doc, "6. 下一步建议", 1)
    add_bullets(
        doc,
        [
            "针对 1_seg03 研究冷启动抗重复特征方法，例如 top-k 多假设初始化 + HMM 延迟决策，而不是单一路径从第一个点开始定死。",
            "把 forward-only 锚点图扩展为“多锚点图集”，在线根据局部似然和方向一致性选择图，而不是先验固定一张图。",
            "设计不依赖真值的完整性指标：候选峰间距、top-k 轨迹分歧、速度可行性、局部磁信息量、HMM 残差。",
            "如果后续采集允许，增加第三天数据，用 4.14 组内调参、5.13 验证、第三天盲测，避免五段数据上的偶然性。",
        ],
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "Siebler, B., Heirich, O., Sand, S. Train Localization with Particle Filter and Magnetic Field Measurements. FUSION 2018. DOI: 10.23919/ICIF.2018.8455298.",
        "Siebler, B., Lehner, A., Sand, S., Hanebeck, U. D. Magnetic Field Mapping of Railway Lines with Graph SLAM. FUSION 2024. DOI: 10.23919/FUSION59988.2024.10706392.",
        "Dieckow, N., Ostaszewski, K., Heinisch, P., Struckmann, H., Ranocha, H. Real-time rail vehicle localisation using spatially resolved magnetic field measurements. arXiv:2507.19327, 2025.",
        "WM-GFM: a novel geomagnetic feature matching positioning method with weak mileage aid. Measurement, 257(D), 118850, 2026. DOI: 10.1016/j.measurement.2025.118850.",
    ]
    add_numbered(doc, refs)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
