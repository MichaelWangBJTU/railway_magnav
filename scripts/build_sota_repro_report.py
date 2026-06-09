from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path.home() / "Desktop" / "磁导航" / "数据" / "codex_railway_magnav"
OUT_ROOT = PROJECT_ROOT / "sota_repro"
OUTPUTS = OUT_ROOT / "outputs"
FIGURES = OUT_ROOT / "figures"
REPORTS = OUT_ROOT / "reports"
DOCX_PATH = REPORTS / "铁路地磁导航SOTA复现与创新点探索报告.docx"


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("铁路地磁导航 SOTA 复现与创新点探索报告")
    set_font(run, 22, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f"数据：良陈铁路小车实验 4.14 / 5.13；生成时间：{datetime.now():%Y-%m-%d %H:%M}")
    set_font(run, 10, color="555555")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, 10.5)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_font(run, 10.5)


def add_prose(doc: Document, text: str, size: float = 10.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size)


def add_formula(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F8FAFC")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("1F3A5F")
    doc.add_paragraph()


def add_method_block(doc: Document, title: str, bullets: list[str], formulas: list[str] | None = None) -> None:
    doc.add_heading(title, level=3)
    add_bullets(doc, bullets)
    for formula in formulas or []:
        add_formula(doc, formula)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_font(run, 9, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_df_table(doc: Document, df: pd.DataFrame, columns: list[tuple[str, str]], title: str | None = None) -> None:
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title)
        set_font(r, 10.5, bold=True, color="1F4D78")
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for cell, (_, header) in zip(table.rows[0].cells, columns):
        shade_cell(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for cell, (col, _) in zip(cells, columns):
            val = row.get(col, "")
            if isinstance(val, float):
                if pd.isna(val):
                    txt = ""
                elif abs(val) >= 100:
                    txt = f"{val:.1f}"
                else:
                    txt = f"{val:.3g}"
            else:
                txt = str(val)
            set_cell_text(cell, txt)
    doc.add_paragraph()


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        p = doc.add_paragraph(f"[缺少图片：{path.name}]")
        set_font(p.runs[0], 10, color="9B1C1C")
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(caption)
    set_font(r, 9, color="555555")


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_font(r, 10.5, bold=True, color="1F3A5F")
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    set_font(r2, 10)
    doc.add_paragraph()


def read_csv(name: str) -> pd.DataFrame:
    path = OUTPUTS / name
    return pd.read_csv(path) if path.exists() else pd.DataFrame()


def main() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    doc.add_heading("1. 核心结论", level=1)
    add_callout(
        doc,
        "一句话结论",
        "在当前良陈铁路约 700 m 数据上，直接照搬 FUSION 2024 铁路磁图 Graph SLAM 并不能复现论文的 1 m 级效果，主要瓶颈是缺少真实轮速里程计、线路较短、磁签名重复导致 false loop closure。更有希望的方向是弱里程辅助的整段/子序列匹配，并显式评估磁特征唯一性。",
    )
    add_bullets(
        doc,
        [
            "窗口级 baseline 中，total + 梯度 NCC 的 150 m 窗口最好，中位绝对误差约 45 m。",
            "弱里程辅助整段对齐把中位绝对误差降到约 29.4 m，说明“相对里程约束 + 磁序列整体匹配”明显优于孤立窗口检索。",
            "辨识度驱动的子序列选择进一步把整体中位误差降到约 18.25 m；在唯一性门限大于 0.20 时，覆盖 50% 片段且中位误差约 1 m。",
            "因此，当前最有价值的创新点不是简单处理方向，而是做“特征唯一性建模 + 弱里程约束 + 置信度门控/拒识 + 鲁棒闭环”。",
        ]
    )

    doc.add_heading("2. 参考文献与方法定位", level=1)
    add_df_table(
        doc,
        pd.DataFrame(
            [
                ["Railway Graph SLAM", "Siebler et al., FUSION 2024", "局部磁图 + 50 m 签名相关，生成一维 pose graph 闭环", "论文报告多趟里程误差约束到约 1 m，loop closure RMSE 约 0.45 m"],
                ["INS + Magnetic Signatures", "Siebler et al., ION GNSS+ 2017", "磁签名匹配提供位置观测，接入 ESKF 抑制 INS 漂移", "铁路磁定位早期强 baseline"],
                ["SLAC", "Siebler et al., NAVIGATION 2023", "Rao-Blackwellized particle filter 同时估计位置和磁强计标定", "适合未标定磁强计和车体磁干扰问题"],
                ["WM-GFM", "Measurement 2026", "弱里程辅助的地磁特征点匹配，降低 DTW 误匹配与复杂度", "平均定位误差报道约 1.78 m"],
                ["Online 1D Magnetic SLAM", "Kok & Solin, MFI 2024", "在线一维磁场 loop closure + 平滑器修正里程漂移", "强调在线闭环检测与决策"],
            ],
            columns=["方向", "代表文献", "核心思想", "对本课题启发"],
        ),
        [("方向", "方向"), ("代表文献", "代表文献"), ("核心思想", "核心思想"), ("对本课题启发", "对本课题启发")],
        "表 1  相关 SOTA / baseline 方法定位",
    )

    p = doc.add_paragraph()
    r = p.add_run("主要来源：")
    set_font(r, 10, bold=True)
    sources = [
        "Magnetic Field Mapping of Railway Lines with Graph SLAM, DOI 10.23919/FUSION59988.2024.10706392, https://elib.dlr.de/204214/",
        "Bounding INS Positioning Errors with Magnetic-Field-Signatures in Railway Environments, DOI 10.33012/2017.15311, https://www.ion.org/publications/pdf.cfm?articleID=15311",
        "Simultaneous Localization and Calibration Methods for a Train-Mounted Magnetometer, DOI 10.33012/navi.557, https://www.ion.org/publications/abstract.cfm?articleID=102996",
        "WM-GFM: a novel geomagnetic feature matching positioning method with weak mileage aid, Measurement 2026, https://www.sciencedirect.com/science/article/pii/S0263224125022092",
        "Online One-Dimensional Magnetic Field SLAM with Loop-Closure Detection, arXiv:2409.01091, https://arxiv.org/abs/2409.01091",
    ]
    add_bullets(doc, sources)
    add_df_table(
        doc,
        pd.DataFrame(
            [
                [
                    "Siebler, Lehner, Sand, Hanebeck",
                    "Magnetic Field Mapping of Railway Lines with Graph SLAM",
                    "Proc. FUSION, 2024",
                    "10.23919/FUSION59988.2024.10706392",
                    "本文复现的主 SOTA：一维铁路 pose graph、磁签名相关闭环、局部磁图思想。",
                ],
                [
                    "Siebler, Heirich, Sand",
                    "Bounding INS Positioning Errors with Magnetic-Field-Signatures in Railway Environments",
                    "ION GNSS+, 2017, pp. 3224-3230",
                    "10.33012/2017.15311",
                    "作为磁签名匹配 + ESKF 的铁路定位 baseline。",
                ],
                [
                    "Siebler, Lehner, Sand, Hanebeck",
                    "Simultaneous Localization and Calibration Methods for a Train-Mounted Magnetometer",
                    "NAVIGATION, 70(1), 2023",
                    "10.33012/navi.557",
                    "启发跨趟/跨日期磁强计偏置、尺度、标定参数与定位联合估计。",
                ],
                [
                    "You, Ji, Wei, Lai, Zhang et al.",
                    "WM-GFM: a novel geomagnetic feature matching positioning method with weak mileage aid",
                    "Measurement, 257, 118850, 2026",
                    "10.1016/j.measurement.2025.118850",
                    "启发弱里程约束、特征点/梯度约束和干扰抑制，不再完全依赖 DTW 弹性匹配。",
                ],
                [
                    "Kok, Solin",
                    "Online One-Dimensional Magnetic Field SLAM with Loop-Closure Detection",
                    "arXiv:2409.01091 / MFI 2024",
                    "arXiv:2409.01091",
                    "启发在线闭环检测、闭环决策和 odometry smoothing。",
                ],
            ],
            columns=["作者", "题名", "出处", "DOI / 链接", "本文使用方式"],
        ),
        [("作者", "作者"), ("题名", "题名"), ("出处", "出处"), ("DOI / 链接", "DOI / 链接"), ("本文使用方式", "本文使用方式")],
        "表 1b  参考文献完整条目与本文借鉴点",
    )

    doc.add_heading("3. 数据与复现实验设置", level=1)
    add_numbered(
        doc,
        [
            "读取 4.14 和 5.13 已处理磁图，包括每 0.5 m 的融合磁图、各趟宽表磁数据和 segment 元数据。",
            "新转换的 SPAN ASCII 日志包含 BESTPOS、BESTVEL、GPGGA、GPRMC、INSPVAX。BESTPOS/BESTVEL 日志头中含 GPS week 与 seconds，可用于后续更严格的里程/速度构建。",
            "由于当前实验没有真实轮速计，Graph SLAM 复现采用 SPAN 位置生成轨道真值，再注入可控尺度误差和随机误差构造伪里程计。这一点与原论文不同，报告中不把它当作完全等价复现。",
            "匹配验证以 4.14 融合磁图为参考图，以 5.13 各 segment 作为待定位曲线，评估预测起点相对 SPAN 对齐坐标的误差。",
        ],
    )
    add_method_block(
        doc,
        "3.1 伪里程计构造方法",
        [
            "目的：FUSION 2024 的 Graph SLAM 输入是轮速里程计 + 磁强计，而当前数据没有真实轮速/编码器。因此我用 SPAN 对齐后的沿轨道距离 s_i 作为真值，再人工注入小的尺度漂移和随机误差，构造一个“可控伪里程计”。这样可以验证磁闭环能否约束里程漂移，但不能把它等同于真实列车轮速计实验。",
            "节点：在 4.14 各有效 segment 上约每 25 m 取一个 Graph SLAM 节点，节点真值为 SPAN 投影到轨道轴上的一维坐标 s_i。",
            "同一连续 segment 内：对相邻节点的真值增量 Δs_i = s_i - s_{i-1} 乘以每趟固定尺度偏差，并叠加小随机噪声，模拟轮速积分误差。",
            "不同 segment 边界：由于实际复现时跨段起停、掉头和数据切分会带来额外不确定性，边界增量加入更大的随机误差。",
            "参数：每个 segment 的尺度偏差 β_seg 服从均值 0.008、标准差 0.003 的正态分布；同段随机噪声标准差 0.15 m；跨段边界随机噪声标准差 0.75 m。最终伪里程计 RMSE 约 2.65 m，最大误差约 5.24 m。",
        ],
        [
            "同段伪里程增量： z_i^odo = (s_i - s_{i-1}) · (1 + β_seg) + ε_i,  β_seg ~ N(0.008, 0.003²),  ε_i ~ N(0, 0.15²)",
            "跨段伪里程增量： z_i^odo = (s_i - s_{i-1}) + η_i,  η_i ~ N(0, 0.75²)",
            "伪里程轨迹： x_0^odo = s_0,  x_i^odo = x_{i-1}^odo + z_i^odo",
        ],
    )
    add_callout(
        doc,
        "注意",
        "伪里程计只是为了复现 Graph SLAM 的优化机制。后续真正做论文实验时，最好使用轮速计、编码器、OBD、视频估速，或者至少用 BESTVEL 积分构造更接近真实系统的弱里程观测。",
    )

    inv = read_csv("span_ascii_inventory.csv")
    if not inv.empty:
        inv_sum = inv.groupby("kind").agg(file_count=("path", "count"), total_rows=("rows", "sum"), total_mb=("size_mb", "sum")).reset_index()
        add_df_table(doc, inv_sum, [("kind", "日志"), ("file_count", "文件数"), ("total_rows", "总行数"), ("total_mb", "总大小 MB")], "表 2  新转换 SPAN ASCII 日志概况")

    doc.add_heading("4. FUSION 2024 Graph SLAM 复现", level=1)
    add_bullets(
        doc,
        [
            "节点间距设为约 25 m，磁签名长度设为 50 m，与论文主要参数保持一致。",
            "相邻节点边来自伪里程计，闭环边来自不同 segment 之间的磁签名相关匹配。",
            "论文中每个节点保存 100 m 局部磁图，并通过滑动相关得到相对位置；本复现保留一维 pose graph 核心思想，但闭环位移估计做了透明简化。",
        ]
    )
    add_method_block(
        doc,
        "4.1 一维 pose graph 模型",
        [
            "状态量：每个节点只估计一个沿轨道方向的位置 x_i，而不是二维 x-y 或三维姿态。这是铁路约束场景的核心简化。",
            "先验边：固定第一个节点的位置，避免整个图平移不定。",
            "里程边：相邻节点之间的相对距离由伪里程计给出。",
            "磁闭环边：如果两个不同 segment 的 50 m 磁签名高度相似，就在这两个节点之间加入一条相对位置约束。",
            "优化器：把所有边写成线性残差，用稀疏最小二乘求解；鲁棒版本对闭环残差做迭代重加权，降低可疑闭环权重。",
        ],
        [
            "状态： X = [x_0, x_1, ..., x_N]^T",
            "先验残差： e_prior = x_0 - s_0",
            "里程残差： e_i^odo = (x_i - x_{i-1}) - z_i^odo",
            "闭环残差： e_ij^lc = (x_i - x_j) - z_ij^lc",
            "目标函数： min_X Σ ||e_prior||²/σ_prior² + Σ ||e_i^odo||²/σ_odo² + Σ ||e_ij^lc||²/σ_lc²",
        ],
    )
    add_method_block(
        doc,
        "4.2 磁闭环检测与本文简化",
        [
            "论文做法：每个节点保存局部磁图，新节点取约 50 m 磁签名，在历史局部磁图中滑动，计算相关系数。相关峰值超过阈值时，峰值位置给出 loop closure 的相对位移。",
            "本文复现：先用 50 m total 磁签名或 total+梯度签名做相关检索，阈值配置包括 0.97 paper-like、0.90 lower threshold 和梯度增强配置。",
            "为了把“闭环检测质量”和“图优化机制”分开诊断：当检测到的闭环在 SPAN 真值上确实接近同一位置时，闭环观测加入约 0.45 m 噪声；当检测是错误闭环时，保留错误约束，以观察 false loop closure 对图优化的破坏。",
            "这也是结果对不上论文的关键：你的短线路里高相关不一定代表同一位置，少数错误闭环会显著拉偏一维图。",
        ],
        [
            "Pearson/NCC 相似度： ρ = Σ[(q_k - q_bar)(m_k - m_bar)] / sqrt(Σ(q_k - q_bar)² Σ(m_k - m_bar)²)",
            "闭环判定： max_d ρ(d) > T_lc；paper-like T_lc = 0.97，本实验还测试了 0.90 和梯度增强。",
        ],
    )
    graph = read_csv("graph_slam_summary.csv")
    if not graph.empty:
        gshow = graph[["method", "node_count", "loop_count", "false_loop_rate_gt2m", "odometry_rmse_m", "slam_rmse_m", "improvement_rmse_pct"]].copy()
        add_df_table(
            doc,
            gshow,
            [
                ("method", "方法"),
                ("node_count", "节点"),
                ("loop_count", "闭环"),
                ("false_loop_rate_gt2m", "错闭环率"),
                ("odometry_rmse_m", "里程 RMSE"),
                ("slam_rmse_m", "SLAM RMSE"),
                ("improvement_rmse_pct", "改善 %"),
            ],
            "表 3  Graph SLAM 复现实验结果",
        )
    add_picture(doc, FIGURES / "graph_slam_summary.png", "图 1  Graph SLAM 复现：当前数据上错误闭环会使优化结果变差。")
    add_callout(
        doc,
        "为什么没有达到论文结果",
        "论文测试段约 1.68 km、四次 run、闭环数量超过 2000，且输入是真实列车里程计；本数据可用轨段约 560 m，节点约 106，闭环只有个位数，且轨道磁特征存在重复。少量 false loop closure 足以把一维图优化拉偏。因此本数据不适合直接照搬高阈值相关闭环作为最终方案。",
    )

    doc.add_heading("5. Baseline 与其他方法复现", level=1)
    add_method_block(
        doc,
        "5.1 窗口级 NCC / 归一化内积匹配",
        [
            "参考图：4.14 融合磁图；待定位曲线：5.13 各 segment 按 0.5 m 重采样后的磁曲线窗口。",
            "候选搜索：把 query 窗口在 4.14 磁图上逐点滑动，计算每个候选起点的相似度，取最高分对应的位置作为预测起点。",
            "特征配置：NCC_total 只用 total；NCC_total_grad 使用 total 和 total 梯度；NCC_total_y_grad 进一步加入 Y anomaly 及其梯度。",
            "归一化：为了减少跨日期偏置和幅值差异，每个窗口先做 robust z-score，即用中位数和 MAD 归一化。多特征版本把多个归一化特征拼接后做内积，因此它是 correlation-like score，不严格限制在 [-1, 1]，表里出现大于 1 的分数是正常的。",
        ],
        [
            "robust z-score： q'_k = (q_k - median(q)) / (1.4826 · MAD(q))",
            "窗口得分： score(d) = mean( concat(q'_features) · concat(m'_{d,features}) )",
            "预测位置： d_hat = argmax_d score(d)",
        ],
    )
    add_method_block(
        doc,
        "5.2 DTW baseline",
        [
            "DTW 用于处理速度不均匀或空间尺度不完全一致的序列。它允许 query 和 map 窗口在局部拉伸/压缩后对齐。",
            "本文只在 50 m 和 100 m 窗口上跑 DTW_total，候选起点每 10 m 搜一次，并使用 Sakoe-Chiba 窄带约束以控制计算量。",
            "结果显示 DTW_total 并没有优于 NCC_total_grad，说明当前误差主要不是速度伸缩，而是磁特征重复和跨日期差异。",
        ],
        [
            "DTW 距离： D(i,j)=|q_i-m_j|+min(D(i-1,j), D(i,j-1), D(i-1,j-1))",
            "预测位置： d_hat = argmin_d DTW(q, m_d)",
        ],
    )
    add_method_block(
        doc,
        "5.3 Viterbi 连续匹配",
        [
            "Viterbi 把位置看作隐状态，把磁场观测作为发射概率，把沿轨道连续运动作为转移概率。",
            "状态是 4.14 磁图上的 0.5 m 网格点；观测似然来自 5.13 total 与参考 total 的 robust z-score 差异；转移模型偏好每步沿距离正向推进约 0.5 m。",
            "这一版 Viterbi 没有取得好结果，主要因为它用的是点级观测，单个磁场点太容易重复；后续应改成窗口似然或特征点似然。",
        ],
        [
            "观测似然： log p(y_t | s_t) = -0.5 · ((q'_t - m'(s_t))/σ_obs)²",
            "转移似然： log p(s_t | s_{t-1}) = -0.5 · ((s_t - s_{t-1} - Δs)/σ_trans)²",
        ],
    )
    base = read_csv("baseline_matching_summary.csv")
    if not base.empty:
        best = base.sort_values(["median_abs_error_m", "rmse_error_m"]).head(8)
        add_df_table(
            doc,
            best,
            [
                ("method", "方法"),
                ("window_m", "窗口 m"),
                ("query_count", "样本数"),
                ("median_abs_error_m", "中位误差 m"),
                ("rmse_error_m", "RMSE m"),
                ("p90_abs_error_m", "P90 m"),
            ],
            "表 4  窗口级匹配 baseline 对比",
        )
    add_picture(doc, FIGURES / "baseline_matching_median_error.png", "图 2  NCC/DTW 等窗口级 baseline 的误差对比。")
    vit = read_csv("continuous_viterbi_summary.csv")
    if not vit.empty:
        add_df_table(
            doc,
            vit,
            [("method", "方法"), ("query_count", "片段数"), ("median_abs_error_m", "中位误差 m"), ("rmse_error_m", "RMSE m")],
            "表 5  连续 Viterbi 匹配结果",
        )

    doc.add_heading("6. 弱里程与辨识度驱动方法", level=1)
    add_method_block(
        doc,
        "6.1 弱里程辅助整段对齐",
        [
            "思想：不把每个短窗口孤立拿去全图检索，而是保留 query 内部的相对里程关系。也就是说，query 曲线内部每个点之间相隔多少米是可信的，只搜索它在参考磁图中的整体绝对起点偏移。",
            "实现：对每个 5.13 segment 取最长不超过 320 m 的中心子段，构造相对距离 r = s - s_start。然后枚举参考图起点 a，比较 q(r) 与 m(a+r)。",
            "特征：测试 total、high-pass total + total 梯度、high-pass total + total 梯度 + high-pass Y anomaly。高通特征用于去掉跨日期慢变偏置，梯度特征强调局部峰谷形状。",
            "结果：WeakMileage_highpass_grad 中位误差约 29.4 m，明显优于固定窗口 NCC。这说明“弱里程约束”能显著减少搜索自由度。",
        ],
        [
            "整体偏移匹配： score(a) = mean_f corr( q_f(r), m_f(a+r) )",
            "预测起点： a_hat = argmax_a score(a)",
            "高通特征： hp_total(s) = total(s) - median_filter(total(s), 30 m)",
        ],
    )
    add_method_block(
        doc,
        "6.2 辨识度驱动子序列选择",
        [
            "动机：同一趟数据中并不是所有片段都适合定位。有些区段磁特征非常唯一，能给出米级定位；有些区段有多个相似候选，强行输出就会出现 100-300 m 级误匹配。",
            "做法：在每个 5.13 segment 内扫描不同长度的子序列，长度集合为 100、140、180、240、320 m，起点每 20 m 试一次。每个候选子序列都进行弱里程匹配，并记录最佳候选和次佳候选之间的分数差。",
            "选择准则：优先选择唯一性更高的子序列，而不是固定取最长或固定取中间段。实际 selection_score = uniqueness margin + 0.25 × best_score。",
            "结果：不门控时中位误差约 18.25 m；当唯一性门限设为 0.20 时，只接受 3/6 个片段，但接受片段的中位误差和 P90 都约 1 m。这说明门控可以把“可定位片段”和“歧义片段”区分开。",
        ],
        [
            "最佳候选： score_best = max_a score(a)",
            "次佳候选： score_second = max_{|a-a_best| >= 20m} score(a)",
            "唯一性指标： U = score_best - score_second",
            "门控规则： 若 U >= τ，则接受本次定位；若 U < τ，则认为磁特征歧义，拒绝输出或继续累积更长序列。",
            "本文关键门限：τ = 0.20 时覆盖率 50%，接受片段中位误差约 1 m。",
        ],
    )
    doc.add_heading("6.3 唯一性门限的含义", level=3)
    add_callout(
        doc,
        "唯一性门限的含义",
        "唯一性门限不是神经网络置信度，也不是概率。它衡量的是匹配得分曲线的峰值是否足够突出：如果第一名和第二名差距很小，说明地图中有多个相似位置，定位结果不可靠；如果差距很大，说明这个磁片段具有较强的空间指纹性，可以更放心地输出位置。",
    )
    weak = read_csv("weak_mileage_alignment_summary.csv")
    if not weak.empty:
        add_df_table(
            doc,
            weak,
            [
                ("method", "方法"),
                ("query_count", "片段数"),
                ("median_abs_error_m", "中位误差 m"),
                ("mean_abs_error_m", "平均误差 m"),
                ("rmse_error_m", "RMSE m"),
                ("p90_abs_error_m", "P90 m"),
            ],
            "表 6  弱里程辅助整段对齐结果",
        )
    distinctive = read_csv("distinctive_subsequence_summary.csv")
    if not distinctive.empty:
        add_df_table(
            doc,
            distinctive,
            [
                ("method", "方法"),
                ("query_count", "片段数"),
                ("median_abs_error_m", "中位误差 m"),
                ("mean_abs_error_m", "平均误差 m"),
                ("rmse_error_m", "RMSE m"),
                ("median_margin", "中位唯一性"),
            ],
            "表 7  辨识度子序列选择结果",
        )
    add_picture(doc, FIGURES / "weak_mileage_alignment_summary.png", "图 3  弱里程辅助整段对齐优于固定窗口检索。")
    add_picture(doc, FIGURES / "confidence_gating_tradeoff.png", "图 4  置信度门控：唯一性门限升高后误差迅速下降，但覆盖率降低。")
    conf = read_csv("confidence_gating_summary.csv")
    if not conf.empty:
        key = conf[(conf["method"] == "DistinctiveSubseq_highpass_grad") & (conf["margin_threshold"].isin([0.0, 0.15, 0.20, 0.25]))]
        add_df_table(
            doc,
            key,
            [
                ("margin_threshold", "唯一性门限"),
                ("accepted_count", "接受数"),
                ("coverage_pct", "覆盖率 %"),
                ("median_abs_error_m", "中位误差 m"),
                ("p90_abs_error_m", "P90 m"),
            ],
            "表 8  辨识度子序列的置信度门控效果",
        )

    doc.add_heading("7. 可作为论文创新点的方向", level=1)
    add_bullets(
        doc,
        [
            "特征唯一性驱动的铁路地磁定位：不再默认每段磁曲线都可定位，而是估计 best/second-best margin、局部熵或峰值尖锐度；高唯一性片段输出位置，低唯一性片段继续累积里程或拒识。",
            "弱里程约束的子序列主动选择：利用车辆相对里程，在已行驶曲线中自动选择最有辨识度的子窗口，而不是固定窗口长度或固定取中间段。当前数据上门限 0.20 时 50% 片段达到约 1 m 中位误差。",
            "鲁棒磁闭环 Graph SLAM：把闭环检测分成候选生成、唯一性打分、几何一致性检验、鲁棒核优化四层，避免短线路相似磁特征造成 false loop closure。",
            "跨日期自标定磁图：借鉴 SLAC，把每趟/每天的 hard-iron、soft-iron、慢变偏置作为隐变量，与位置匹配联合估计。这样比简单的方向修正更贴合当前 4.14/5.13 三轴幅值差异。",
            "不确定性磁图：每个地图点不仅存均值/中位数，还存跨趟方差、MAD、特征唯一性和可定位性概率。定位时用地图不确定性加权，而不是所有区段等权匹配。",
        ]
    )

    doc.add_heading("8. 下一步实验建议", level=1)
    add_numbered(
        doc,
        [
            "把 BESTVEL 积分与 SPAN 位置差分对比，确认能否构造真实伪里程计，而不是人工注入漂移。",
            "补采或整理轮速/编码器数据；如果没有轮速，至少记录小车控制速度或使用视频/时间估计相对里程。",
            "在 4.14 内部做留一趟验证，再做 4.14 到 5.13 跨日期验证，分离“同日可重复性”和“跨日变化”两类问题。",
            "实现 SLAC 风格的每趟偏置/尺度联合估计，重点验证三轴分量是否能从当前不可用变成有效辅助特征。",
            "把辨识度门控方法接入在线定位：低置信区不报位置，高置信区报绝对位置，并给出置信度。",
        ]
    )

    doc.add_heading("附录：输出文件", level=1)
    add_df_table(
        doc,
        pd.DataFrame(
            [
                ["复现实验代码", str(OUT_ROOT / "code")],
                ["结果表", str(OUTPUTS)],
                ["关键图片", str(FIGURES)],
                ["本报告", str(DOCX_PATH)],
            ],
            columns=["类别", "路径"],
        ),
        [("类别", "类别"), ("路径", "路径")],
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
